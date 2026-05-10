#!/usr/bin/env python3
"""
Generates backend/templates/cover_template.docx — the cover page.
Run from the backend/ directory:
    python3 create_cover_template.py
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BORDER  = 'B8CCE4'  # light blue matching original
N       = 14        # total columns
CS      = 1         # content start index
CE      = 12        # content end index  (cols 1-12 merged)

# ── helpers ──────────────────────────────────────────────────────────────────

def _borders(cell, show=True):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    b = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        if show:
            el.set(qn('w:val'),   'single')
            el.set(qn('w:sz'),    '4')
            el.set(qn('w:color'), BORDER)
        else:
            el.set(qn('w:val'),   'none')
            el.set(qn('w:sz'),    '0')
            el.set(qn('w:color'), 'FFFFFF')
        el.set(qn('w:space'), '0')
        b.append(el)
    tcPr.append(b)


def _row_height(row, twips, rule='exact'):
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    for old in trPr.findall(qn('w:trHeight')):
        trPr.remove(old)
    h = OxmlElement('w:trHeight')
    h.set(qn('w:val'),   str(twips))
    h.set(qn('w:hRule'), rule)
    trPr.append(h)


def spacer(tbl, h=95):
    """Visible grid spacer row."""
    row = tbl.add_row()
    _row_height(row, h)
    for cell in row.cells:
        _borders(cell, True)
        for p in cell.paragraphs:
            p.clear()
    return row


def crow(tbl, min_h=None):
    """
    Content row: col-0 and col-13 stay as narrow grid columns;
    cols 1-12 are merged into one content area.
    Returns (row, content_cell).
    """
    row = tbl.add_row()
    if min_h:
        _row_height(row, min_h, 'atLeast')

    # Merge cols 1 … 12
    for _ in range(CE - CS):
        row.cells[CS].merge(row.cells[CS + 1])

    _borders(row.cells[0],   True)   # left edge
    _borders(row.cells[CS],  False)  # content (no inner borders)
    _borders(row.cells[-1],  True)   # right edge

    cc = row.cells[CS]
    while len(cc.paragraphs) > 1:
        cc.paragraphs[-1]._element.getparent().remove(
            cc.paragraphs[-1]._element)
    cc.paragraphs[0].clear()
    return row, cc


def p(cell, text, bold=False, ul=False, sz=11, sb=0, sa=4):
    """Add a centred paragraph with one run to a cell."""
    para = cell.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(sb)
    para.paragraph_format.space_after  = Pt(sa)
    run = para.add_run(text)
    run.bold      = bold
    run.underline = ul
    run.font.size = Pt(sz)
    return para


def ctrl(cell, tag):
    """Invisible docxtpl control-tag paragraph."""
    para = cell.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run = para.add_run(tag)
    run.font.size = Pt(0.5)
    return para


# ── builder ──────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    for s in doc.sections:
        s.page_width    = Mm(210)
        s.page_height   = Mm(297)
        s.top_margin    = Cm(1.0)
        s.bottom_margin = Cm(1.0)
        s.left_margin   = Cm(1.5)
        s.right_margin  = Cm(1.5)

    # Remove default empty paragraph
    for para in list(doc.paragraphs):
        para._element.getparent().remove(para._element)

    # Table  (14 cols: 0.4 | 12×1.44 | 0.4 = 18 cm usable)
    tbl = doc.add_table(rows=0, cols=N)
    tbl.style = 'Table Grid'

    # ── top spacer ──────────────────────────────────────────────────────────
    spacer(tbl)

    # ── Title ───────────────────────────────────────────────────────────────
    _, c = crow(tbl, 400)
    p(c, 'Property Valuation Report', bold=True, sz=18, sb=10, sa=2)
    p(c, '({{ property_mortgaged }})',  bold=True, sz=15, sb=0,  sa=10)

    spacer(tbl)

    # ── Client ──────────────────────────────────────────────────────────────
    _, c = crow(tbl, 460)
    p(c, 'Client',                        bold=True, ul=True, sz=13, sb=8, sa=4)
    p(c, '{{ client_full_name }}',         bold=True,          sz=13, sb=0, sa=4)
    p(c, '{{ client_vdc_ward }}',                              sz=11, sb=0, sa=2)
    p(c, '{{ client_district_country }}',                      sz=11, sb=0, sa=2)
    p(c, 'Contact No: {{ client_phone }}',                     sz=11, sb=0, sa=8)

    spacer(tbl)

    # ── Properties to be Valued ─────────────────────────────────────────────
    _, c = crow(tbl, 480)
    p(c, 'Properties to be Valued', bold=True, ul=True, sz=13, sb=8, sa=4)
    p(c, 'Land & Building Located at', bold=True,        sz=11, sb=0, sa=4)
    p(c, 'Plot No:  {{ plot_no }}',                      sz=11, sb=0, sa=2)
    p(c, '{{ property_location_full }}',                  sz=11, sb=0, sa=2)
    p(c, '{{ sabik_display }}',                           sz=11, sb=0, sa=8)

    spacer(tbl)

    # ── Owner 1 ─────────────────────────────────────────────────────────────
    _, c = crow(tbl, 460)
    p(c, 'Owner 1',                        bold=True, ul=True, sz=13, sb=8, sa=4)
    p(c, '{{ owner1_name }}',               bold=True,          sz=13, sb=0, sa=4)
    p(c, '{{ owner1_vdc_ward }}',                               sz=11, sb=0, sa=2)
    p(c, '{{ owner1_district_country }}',                       sz=11, sb=0, sa=2)
    p(c, 'Contact No: {{ owner1_phone }}',                      sz=11, sb=0, sa=8)

    spacer(tbl)

    # ── Owner 2 (conditional — shown only when a second owner exists) ────────
    _, c = crow(tbl, 460)
    ctrl(c, "{%p if owner2_name and owner2_name != '—' %}")
    p(c, 'Owner 2',                        bold=True, ul=True, sz=13, sb=8, sa=4)
    p(c, '{{ owner2_name }}',               bold=True,          sz=13, sb=0, sa=4)
    p(c, '{{ owner2_vdc_ward }}',                               sz=11, sb=0, sa=2)
    p(c, '{{ owner2_district_country }}',                       sz=11, sb=0, sa=2)
    p(c, 'Contact No: {{ owner2_phone }}',                      sz=11, sb=0, sa=8)
    ctrl(c, '{%p endif %}')

    spacer(tbl)

    # ── Submitted to ────────────────────────────────────────────────────────
    _, c = crow(tbl, 360)
    p(c, 'Submitted to',   bold=True, ul=True, sz=13, sb=8, sa=8)
    p(c, '{{ bank_name }}', bold=True, ul=True, sz=13, sb=0, sa=4)
    p(c, '{{ bank_branch }}',                   sz=11, sb=0, sa=8)

    spacer(tbl)

    # ── Fair Market Value ────────────────────────────────────────────────────
    _, c = crow(tbl, 180)
    p(c, 'Fair Market Value of the Land & Building: NRs. {{ fair_market_value }}',
      bold=True, sz=12, sb=8, sa=8)

    spacer(tbl)

    # ── Prepared by ─────────────────────────────────────────────────────────
    _, c = crow(tbl, 160)
    p(c, 'Prepared by', bold=True, ul=True, sz=13, sb=8, sa=8)

    # Logo placeholder row – _inject_logo_post fills this space
    _, c = crow(tbl, 1500)
    p(c, '', sz=11, sb=0, sa=0)

    spacer(tbl)

    # ── Firm details ────────────────────────────────────────────────────────
    _, c = crow(tbl, 560)
    p(c, '{{ firm_name }}',        sz=11, sb=8, sa=2)
    p(c, '{{ firm_street }}',      sz=11, sb=0, sa=2)
    p(c, '{{ firm_city_country }}', sz=11, sb=0, sa=2)
    p(c, 'Tel: {{ firm_phone }}',  sz=11, sb=0, sa=2)
    p(c, 'Email: {{ firm_email }}', sz=11, sb=0, sa=8)

    spacer(tbl)
    spacer(tbl)

    out = os.path.join(os.path.dirname(__file__), 'templates', 'cover_template.docx')
    doc.save(out)
    print(f'✓  cover_template.docx  →  {out}')


if __name__ == '__main__':
    build()
