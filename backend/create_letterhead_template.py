#!/usr/bin/env python3
"""
Creates backend/templates/letterhead_template.docx — property schedule table.
Run from backend/ directory:
    python3 create_letterhead_template.py
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BORDER  = 'B8CCE4'   # light blue border (matches cover)
HDR_BG  = 'D9E1F2'   # header cell background
AC_BG   = 'BFBFBF'   # "Area Considered" shaded row

# 10 columns: SN | Owner | Location | PlotNo | Lalpurja | Meas | Amount | Status | Sqft | BldgAmt
# 18 cm usable width (210mm - 1.5cm*2 margins)
COLS  = 10
COL_W = [0.9, 2.8, 3.2, 1.2, 1.9, 2.0, 1.8, 1.8, 1.4, 1.0]  # sum = 18.0 cm


# ── XML helpers ────────────────────────────────────────────────────────────────

def _borders(cell, color=BORDER, sz='6'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    b = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    sz)
        el.set(qn('w:color'), color)
        el.set(qn('w:space'), '0')
        b.append(el)
    tcPr.append(b)


def _shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)


def _valign(cell, val='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:vAlign')):
        tcPr.remove(old)
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), val)
    tcPr.append(va)


def _vmerge_start(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vm = OxmlElement('w:vMerge')
    vm.set(qn('w:val'), 'restart')
    tcPr.append(vm)


def _vmerge_cont(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vm = OxmlElement('w:vMerge')
    tcPr.append(vm)
    # Clear text so continuation cells are truly empty
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)
    cell.paragraphs[0].clear()


def _row_height(row, twips, rule='exact'):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    for old in trPr.findall(qn('w:trHeight')):
        trPr.remove(old)
    h = OxmlElement('w:trHeight')
    h.set(qn('w:val'),   str(twips))
    h.set(qn('w:hRule'), rule)
    trPr.append(h)


def _write(cell, text, bold=False, sz=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, italic=False):
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.font.size = Pt(sz)


# ── Builder ────────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    for s in doc.sections:
        s.page_width    = Mm(210)
        s.page_height   = Mm(297)
        s.top_margin    = Cm(1.8)
        s.bottom_margin = Cm(1.8)
        s.left_margin   = Cm(1.5)
        s.right_margin  = Cm(1.5)

    for para in list(doc.paragraphs):
        para._element.getparent().remove(para._element)

    tbl = doc.add_table(rows=0, cols=COLS)
    tbl.style = 'Table Grid'

    # Column widths
    tblGrid = OxmlElement('w:tblGrid')
    for w in COL_W:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(round(w * 1440 / 2.54))))
        tblGrid.append(gc)
    tbl._tbl.insert(0, tblGrid)

    def new_row(h=None, rule='atLeast'):
        r = tbl.add_row()
        if h:
            _row_height(r, h, rule)
        for cell in r.cells:
            _borders(cell)
        return r

    # ── Row 0: Main group headers (S.N./Owner/Location/Plot span 2 rows; Land/Building span 3 cols each) ──
    r = new_row(500, 'atLeast')
    c = r.cells

    for i, hdr in enumerate(['S.N.', 'Owner of the\nproperty', 'Location of\nProperty', 'Plot\nNo.']):
        _write(c[i], hdr, bold=True, sz=9.5)
        _shading(c[i], HDR_BG); _valign(c[i]); _vmerge_start(c[i])

    # "Details of Land" spans cols 4–6
    land_cell = c[4]
    c[4].merge(c[6])
    _write(land_cell, 'Details of Land (Area in Square Meter)', bold=True, sz=9.5)
    _shading(land_cell, HDR_BG); _valign(land_cell)

    # "Description of the Building" spans cols 7–9
    bldg_cell = c[7]
    c[7].merge(c[9])
    _write(bldg_cell, 'Description of the Building', bold=True, sz=9.5)
    _shading(bldg_cell, HDR_BG); _valign(bldg_cell)

    # ── Row 1: Sub-headers ────────────────────────────────────────────────────
    r = new_row(480, 'atLeast')
    c = r.cells

    # Cols 0–3: vertical merge continuation
    for i in range(4):
        _vmerge_cont(c[i]); _shading(c[i], HDR_BG); _valign(c[i])

    for i, hdr in enumerate(['Area as per\nLalpurja', 'Area as per\nMeasurement',
                              'Amount', 'Status', 'Area\n(Sq.ft.)', 'Amount'], start=4):
        _write(c[i], hdr, bold=True, sz=9)
        _shading(c[i], HDR_BG); _valign(c[i])

    # ── Row 2: Data row — Lalpurja area ───────────────────────────────────────
    r = new_row(380, 'atLeast')
    c = r.cells

    _write(c[0], '1', sz=9.5);                               _valign(c[0]); _vmerge_start(c[0])
    _write(c[1], '{{ owners_plain }}', sz=9);                 _valign(c[1]); _vmerge_start(c[1])
    _write(c[2], '{{ property_location_full }}', sz=9);       _valign(c[2]); _vmerge_start(c[2])
    _write(c[3], '{{ plot_no }}', sz=9.5);                    _valign(c[3]); _vmerge_start(c[3])
    _write(c[4], '{{ land_area_lorc }}', sz=9.5)
    _write(c[5], '{{ land_area_measured }}', sz=9.5)
    _write(c[6], '{{ fair_market_value_raw }}', sz=9.5);      _valign(c[6]); _vmerge_start(c[6])
    _write(c[7], '{{ building_status }}', sz=9.5);            _valign(c[7]); _vmerge_start(c[7])
    _write(c[8], '{{ building_area_sqft }}', sz=9.5);         _valign(c[8]); _vmerge_start(c[8])
    _write(c[9], '{{ building_fmv }}', sz=9.5);               _valign(c[9]); _vmerge_start(c[9])

    # ── Row 3: "Area Considered" shaded label ─────────────────────────────────
    r = new_row(360, 'exact')
    c = r.cells

    for i in (0, 1, 2, 3): _vmerge_cont(c[i])

    ac_cell = c[4]
    c[4].merge(c[5])
    _write(ac_cell, 'Area Considered', bold=True, italic=True, sz=9.5)
    _shading(ac_cell, AC_BG); _valign(ac_cell)

    for i in (6, 7, 8, 9): _vmerge_cont(c[i])

    # ── Row 4: Considered area value ──────────────────────────────────────────
    r = new_row(380, 'atLeast')
    c = r.cells

    for i in (0, 1, 2, 3): _vmerge_cont(c[i])

    consid_cell = c[4]
    c[4].merge(c[5])
    _write(consid_cell, '{{ land_area_considered }}', bold=True, sz=9.5)
    _valign(consid_cell)

    for i in (6, 7, 8, 9): _vmerge_cont(c[i])

    # ── Row 5: Sheet No. + Building dimension headers ─────────────────────────
    r = new_row(360, 'atLeast')
    c = r.cells

    sheet_cell = c[0]
    c[0].merge(c[6])
    _write(sheet_cell, 'Sheet No: {{ sheet_no }}', sz=9.5); _valign(sheet_cell)

    _write(c[7], 'Length',  bold=True, sz=9.5); _valign(c[7]); _shading(c[7], HDR_BG)
    _write(c[8], 'Breadth', bold=True, sz=9.5); _valign(c[8]); _shading(c[8], HDR_BG)
    _write(c[9], 'Height',  bold=True, sz=9.5); _valign(c[9]); _shading(c[9], HDR_BG)

    # ── Row 6: Building dimension values ──────────────────────────────────────
    r = new_row(360, 'atLeast')
    c = r.cells

    dim_left = c[0]
    c[0].merge(c[6])
    _write(dim_left, '', sz=9.5); _valign(dim_left)

    _write(c[7], '{{ building_length }}',  sz=9.5); _valign(c[7])
    _write(c[8], '{{ building_breadth }}', sz=9.5); _valign(c[8])
    _write(c[9], '{{ building_height }}',  sz=9.5); _valign(c[9])

    # ── Row 7: Total ──────────────────────────────────────────────────────────
    r = new_row(420, 'atLeast')
    c = r.cells

    total_label = c[0]
    total_val   = c[7]
    c[0].merge(c[6])
    c[7].merge(c[9])

    _write(total_label, 'Total Fair Market Value of Land & Building Project',
           bold=True, sz=9.5); _valign(total_label); _shading(total_label, HDR_BG)
    _write(total_val, '{{ fair_market_value }}',
           bold=True, sz=9.5, align=WD_ALIGN_PARAGRAPH.RIGHT); _valign(total_val)

    out = os.path.join(os.path.dirname(__file__), 'templates', 'letterhead_template.docx')
    doc.save(out)
    print(f'✓  letterhead_template.docx  →  {out}')


if __name__ == '__main__':
    build()
