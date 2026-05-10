#!/usr/bin/env python3
"""
Fix paragraph alignment in proposal_template.docx:
  - Cover page content (before first page/section break): CENTER
  - Body paragraphs outside tables: LEFT
  - Table cell content: CENTER
Run from backend/: python3 fix_proposal_alignment.py
"""
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

TEMPLATES = os.path.join(os.path.dirname(__file__), 'templates')


def _has_page_break(para):
    for run in para.runs:
        for br in run._element.findall(qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                return True
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
        return True
    return False


def fix():
    path = os.path.join(TEMPLATES, 'proposal_template.docx')
    doc = Document(path)

    # Fix top-level (non-table) paragraphs
    cover_done = False
    for para in doc.paragraphs:
        if cover_done:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if _has_page_break(para):
            cover_done = True

    # Fix all table cell content to CENTER
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(path)
    print(f'✓  Alignment fixed  →  {path}')


if __name__ == '__main__':
    fix()
