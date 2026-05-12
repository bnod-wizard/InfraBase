"""
Document Service - Merges DB data into .docx templates using docxtpl.

Logo injection:
  Place the firm logo at  backend/templates/firm_logo.png
  (PNG or JPG, ideally ~400 px wide, transparent background).
  It will be inserted automatically after the "Prepared by" line in
  any generated cover or proposal document.
"""
import os, io
from docxtpl import DocxTemplate

TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'templates')

TEMPLATE_FILES = {
    'letterhead': 'letterhead_template.docx',
    'cover':      'cover_template.docx',
    'proposal':   'proposal_template.docx',
}


# ── Formatting helpers ─────────────────────────────────────────────────────────

def _v(value, fallback='—'):
    if value is None or value == '':
        return fallback
    return str(value)

def _money(value):
    if value is None:
        return '—'
    try:
        f = float(value)
        return f'{f:,.2f}'
    except (ValueError, TypeError):
        return '—'

def _yn(flag, yes_text='Yes', no_text='No'):
    if flag is True or str(flag).lower() in ('yes','true','1'):
        return yes_text
    return no_text

def _to_float(value, fallback=0.0):
    try:
        return float(value)
    except (ValueError, TypeError):
        return fallback


# ── Parties helpers ────────────────────────────────────────────────────────────

def _client_full_name(c):
    et = _v(c.get('entity_type'), 'individual')
    if et == 'company':
        return _v(c.get('first_name'), 'Unknown Company')
    title = _v(c.get('title'), '')
    name  = f'{c.get("first_name","") or ""} {c.get("last_name","") or ""}'.strip()
    return f'{title} {name}'.strip()

def _owner_full_name(o):
    title = _v(o.get('title'), '')
    return f'{title} {_v(o.get("owner_name",""), "")}'.strip()

def _party_address(ward, vdc, district, address):
    parts = [p for p in [
        ward and f'Ward No. {ward}',
        _v(vdc, ''),
        _v(district, ''),
    ] if p]
    if parts:
        return ', '.join(parts)
    return _v(address, '—')

def _owners_plain(owners):
    """E.g. 'Mr. A & Mrs. B'"""
    return ' & '.join(_owner_full_name(o) for o in owners) if owners else '—'

def _owners_block(owners):
    """Multi-line block for cover: Owner 1 / Owner 2 …"""
    lines = []
    for i, o in enumerate(owners, 1):
        name  = _owner_full_name(o)
        addr  = _party_address(o.get('ward_no'), o.get('vdc_municipality'),
                               o.get('district'), o.get('address'))
        dist  = _v(o.get('district'), '')
        lines.append(f'Owner {i}\n{name}\n{addr}\n{dist} District, Nepal\nContact No: {_v(o.get("phone") or o.get("mobile"),"—")}')
    return '\n\n'.join(lines)

def _owners_detail_block(owners):
    """Detailed block for Synopsis section."""
    blocks = []
    for i, o in enumerate(owners, 1):
        name = _owner_full_name(o)
        ward = _v(o.get('ward_no'), '—')
        vdc  = _v(o.get('vdc_municipality'), '—')
        dist = _v(o.get('district'), '—')
        cit  = _v(o.get('citizenship_no'), '—')
        cit_date   = _v(o.get('citizenship_issued_date'), '—')
        cit_office = _v(o.get('citizenship_issued_office'), '—')
        father     = _v(o.get('father_name'), '—')
        grand      = _v(o.get('grandfather_name') or o.get('husband_name'), '—')
        grand_label = "Husband's Name" if o.get('husband_name') else "Grand Father's Name"
        phone = _v(o.get('phone') or o.get('mobile'), '—')

        block = (
            f'Owner {i}          : {name}\n'
            f'Ward No.           : {ward}\n'
            f'Municipality/VDC   : {vdc}\n'
            f'District           : {dist}\n'
            f'Citizenship No.    : {cit}\n'
            f'Issued Date        : {cit_date}\n'
            f'Issued Office      : {cit_office}\n'
            f"Father's Name      : {father}\n"
            f'{grand_label:<19}: {grand}\n'
            f'Contact No.        : {phone}'
        )
        blocks.append(block)
    return '\n\n'.join(blocks)

def _client_entity_label(c):
    """Returns e.g. 'Individual (Male)' or 'Company'."""
    et = _v(c.get('entity_type'), 'individual').capitalize()
    if et.lower() == 'individual':
        gender = _v(c.get('gender'), '').capitalize()
        return f'Individual ({gender})' if gender else 'Individual'
    return et

def _client_detail_block(c):
    """Formatted block matching the owners_detail_block style."""
    name   = _client_full_name(c)
    ward   = _v(c.get('ward_no'), '—')
    vdc    = _v(c.get('vdc_municipality'), '—')
    dist   = _v(c.get('district'), '—')
    cit    = _v(c.get('citizenship_no'), '—')
    cit_date   = _v(c.get('citizenship_issued_date'), '—')
    cit_office = _v(c.get('citizenship_issued_office'), '—')
    father     = _v(c.get('father_name'), '—')
    grand      = _v(c.get('grandfather_name') or c.get('husband_name'), '—')
    grand_label = "Husband's Name" if c.get('husband_name') else "Grand Father's Name"
    phone  = _v(c.get('phone') or c.get('mobile'), '—')
    return (
        f'Client             : {name}\n'
        f'Ward No.           : {ward}\n'
        f'Municipality/VDC   : {vdc}\n'
        f'District           : {dist}\n'
        f'Citizenship No.    : {cit}\n'
        f'Issued Date        : {cit_date}\n'
        f'Issued Office      : {cit_office}\n'
        f"Father's Name      : {father}\n"
        f'{grand_label:<19}: {grand}\n'
        f'Contact No.        : {phone}'
    )

def _road_access(prop):
    field = _v(prop.get('road_access_field'), '')
    if field and field != '—':
        return field
    width = _v(prop.get('road_width'), '')
    rtype = _v(prop.get('road_type'), '')
    side  = _v(prop.get('road_side'), '')
    parts = [p for p in [width and f'{width}ft wide', rtype, side and f'at {side} Side'] if p]
    return ' '.join(parts) if parts else '—'

def _traditional(val, fallback='0'):
    """Split a traditional R-A-P-D string like '0-3-0-0' or use individual fields."""
    if isinstance(val, str) and '-' in val:
        parts = val.split('-')
        if len(parts) == 4:
            return parts
    return [fallback, fallback, fallback, fallback]

def _merit_lines(prop, val):
    """Generate importance bullet points from property data."""
    merits = []
    vdc  = _v(prop.get('vdc_municipality'), '')
    ward = _v(prop.get('ward_no'), '')
    loc  = f'{vdc}-{ward}' if ward else vdc
    if loc:
        merits.append(f'The considered site is located at {loc}.')
    road = _road_access(prop)
    if road and road != '—':
        merits.append(f'Access Road to the considered site is from {road}.')
    dist = _v(prop.get('public_transport_distance') or prop.get('nearest_landmark'), '')
    if dist:
        merits.append(f'Public Transportation: {dist}.')
    services = []
    for s, label in [('water_supply','Water Supply Line'),('sewerage','Sewerage Pipe Line'),
                     ('electricity_line','Electricity Line'),('telephone','Telephone Line'),
                     ('tv_cable','TV Cable')]:
        if _yn(prop.get(s)) == 'Yes':
            services.append(label)
    if services:
        merits.append(f'Physical infrastructure services such as {", ".join(services)} are available.')
    bank = _v(val.get('bank_name') if val else None, '')
    if bank and prop.get('nearest_landmark'):
        merits.append(f'Approx. {_v(prop.get("public_transport_distance"),"—")} from {prop.get("nearest_landmark","the bank")}.')
    # Pad to 5
    while len(merits) < 5:
        merits.append('—')
    return merits[:5]


# ── Main context builder ───────────────────────────────────────────────────────

def build_context(hierarchy, valuation):
    account    = hierarchy.get('account', {})
    clients    = hierarchy.get('clients', [])
    owners     = hierarchy.get('owners', [])
    properties = hierarchy.get('properties', [])
    prop       = properties[0] if properties else {}
    val        = valuation or {}

    # Firm info – prefer valuation metadata, fall back to account
    firm_name    = _v(val.get('firm_name'),    account.get('account_name', ''))
    firm_address = _v(val.get('firm_address'), account.get('address', ''))
    firm_phone   = _v(val.get('firm_phone'),   account.get('phone', ''))
    firm_email   = _v(val.get('firm_email'),   account.get('email', ''))

    # First client
    c = clients[0] if clients else {}
    client_name   = _client_full_name(c) if c else '—'
    c_ward   = _v(c.get('ward_no'), '')
    c_vdc    = _v(c.get('vdc_municipality'), '')
    c_dist   = _v(c.get('district'), '')
    client_address_line = f'Ward No. {c_ward}' if c_ward else _v(c.get('address'), '—')
    client_address_full = _party_address(c.get('ward_no'), c.get('vdc_municipality'),
                                         c.get('district'), c.get('address'))
    client_district_country = f'{c_dist} District, Nepal' if c_dist else _v(c.get('country'), '')

    # Property location string
    p_ward      = _v(prop.get('ward_no'), '')
    p_vdc       = _v(prop.get('vdc_municipality'), '')
    p_dist      = _v(prop.get('district'), '')
    sabik_vdc   = _v(prop.get('sabik_vdc'), '')
    sabik_ward  = _v(prop.get('sabik_ward_no'), '')
    # Build sabik string: prefer separate fields, fall back to legacy combined `sabik`
    if sabik_vdc or sabik_ward:
        sabik_str = f'{sabik_vdc}-{sabik_ward}' if sabik_vdc and sabik_ward else (sabik_vdc or sabik_ward)
    else:
        sabik_str = _v(prop.get('sabik'), '')
    present_part = f'{p_vdc}-{p_ward}' if p_ward else p_vdc
    prop_loc_parts = [p for p in [present_part, sabik_str and f'(Sabik {sabik_str})', p_dist and f'{p_dist} District'] if p]
    property_location_full = ', '.join(prop_loc_parts) or _v(prop.get('address'), '—')

    # Land area traditional
    lorc_parts   = _traditional(prop.get('land_area_lorc_trad'))
    meas_parts   = _traditional(prop.get('land_area_meas_trad'))
    ded_trad     = prop.get('land_area_ded_trad') or prop.get('land_area_meas_trad')
    deduct_parts = _traditional(ded_trad)
    # Area considered = after deduction; fall back to deduction then measurement
    consid_trad  = prop.get('land_area_ded_trad') or prop.get('land_area_meas_trad')
    consid_parts = _traditional(consid_trad)

    # Influencing factors
    bool_fields = {
        'near_army_barracks':    'near_army_barracks_text',
        'near_cremation_area':   'near_cremation_area_text',
        'near_dumping_site':     'near_dumping_site_text',
        'near_fuel_depot':       'near_fuel_depot_text',
        'near_hazardous_factory':'near_hazardous_factory_text',
        'near_high_tension_line':'near_high_tension_line_text',
        'near_monument':         'near_monument_text',
        'near_river_stream':     'near_river_stream_text',
        'near_temple':           'near_temple_text',
        'water_logging':         'water_logging_text',
    }

    merits = _merit_lines(prop, val)

    # Importance text
    bank_name = _v(val.get('bank_name'), '—')
    land_use  = _v(prop.get('notes') or prop.get('address'), 'residential purposes')
    importance_text = (
        f'The Property is located at {property_location_full}. '
        f'The site can be accessed from {bank_name} and travelling around '
        f'{_v(prop.get("public_transport_distance") or prop.get("nearest_landmark"), "the area")}.'
    )
    accessibility_text = (
        f'The said site is linked via {_road_access(prop)}. '
        f'The site can be accessed from {bank_name} and travelling around '
        f'{_v(prop.get("public_transport_distance") or prop.get("nearest_landmark"), "the area")}, '
        f'We can reach this Property.'
    )

    # Land area in sq.ft (1 sqm = 10.7639 sqft)
    lorc_sqft = round(_to_float(prop.get('land_area_lorc') or 0) * 10.7639, 2)
    meas_sqft = round(_to_float(prop.get('land_area_measured') or 0) * 10.7639, 2)
    ded_sqm   = _to_float(prop.get('land_area_deducted') or prop.get('land_area_measured') or 0)
    ded_sqft  = round(ded_sqm * 10.7639, 2)
    consid_sqm = _to_float(prop.get('considered_area') or prop.get('land_area_deducted') or prop.get('land_area_measured') or 0)
    consid_sqft = round(consid_sqm * 10.7639, 2)

    # Triangle measurement totals — read from dynamic triangle arrays
    def _tri_total_sqft(area_obj):
        tris = (area_obj or {}).get('triangles', [])
        if tris:
            return sum(_to_float(t.get('area_sqft') or 0) for t in tris)
        # stored total fallback
        stored = _to_float((area_obj or {}).get('total_sqft') or 0)
        return stored

    lm_obj = prop.get('land_area_as_per_measurement') or {}
    lm_total_sqft_raw = _tri_total_sqft(lm_obj) or meas_sqft
    lm_total_sqft = str(round(lm_total_sqft_raw, 2)) if lm_total_sqft_raw else '—'
    lm_total_sqm  = str(round(lm_total_sqft_raw / 10.7639, 2)) if lm_total_sqft_raw else '—'
    lm_triangles  = (lm_obj.get('triangles') or [])

    ded_obj = prop.get('land_area_after_deduction') or {}
    ded_total_sqft_raw = _tri_total_sqft(ded_obj) or ded_sqft
    ded_total_sqft = str(round(ded_total_sqft_raw, 2)) if ded_total_sqft_raw else '—'
    ded_total_sqm  = str(round(ded_total_sqft_raw / 10.7639, 2)) if ded_total_sqft_raw else '—'
    ded_triangles  = (ded_obj.get('triangles') or [])

    road_ded_pct   = _to_float(prop.get('road_deduction_percent') or 0)
    road_ded_amt   = round(ded_total_sqft_raw * road_ded_pct / 100, 2) if ded_total_sqft_raw else 0
    ded_after_sqft_raw = round(ded_total_sqft_raw - road_ded_amt, 2) if ded_total_sqft_raw else 0
    ded_after_sqft = str(ded_after_sqft_raw) if ded_after_sqft_raw else '—'
    ded_after_sqm  = str(round(ded_after_sqft_raw / 10.7639, 2)) if ded_after_sqft_raw else '—'

    lorc_obj = prop.get('land_area_as_per_lalpurja') or {}
    lorc_triangles = (lorc_obj.get('triangles') or [])

    # Lalpurja area in Sq.m. and decimal Aana (from R-A-P-D)
    lorc_sqm  = str(round(lorc_sqft / 10.7639, 2)) if lorc_sqft else '—'
    lorc_parts_raw = _traditional(prop.get('land_area_lorc_trad'))
    try:
        lr, la, lp, ld = (float(x) for x in lorc_parts_raw)
        lorc_aana_dec = round(lr * 16 + la + lp / 4 + ld / 16, 2)
        lorc_aana = str(lorc_aana_dec) if lorc_aana_dec else '—'
    except Exception:
        lorc_aana = '—'

    # Did area decrease from LORC to measurement?
    lorc_num = _to_float(prop.get('land_area_measured') or 0)
    meas_num = _to_float(prop.get('land_area_lorc') or 0)
    area_change = 'Decreased' if lorc_num < meas_num else 'Increased'

    # Summary remarks: prefer stored field, fall back to computed merit lines
    stored_remarks = _v(prop.get('summary_remarks'), '')
    if stored_remarks and stored_remarks != '—':
        summary_remarks = stored_remarks
    else:
        summary_remarks_lines = [merits[i] for i in range(min(5, len(merits))) if merits[i] != '—']
        summary_remarks = '\n'.join(summary_remarks_lines) if summary_remarks_lines else '—'

    # FMV/Distress
    fmv_land     = _money(prop.get('fair_market_value_land'))
    fmv_building = _money(prop.get('fair_market_value_building'))
    fmv_total    = _money(prop.get('fair_market_value_total'))
    dv_total     = _money(prop.get('distress_value_total'))

    # Avg FMV rate per Aana
    govt_rate = _to_float(prop.get('government_rate_per_aana') or 0)
    comm_rate = _to_float(prop.get('commercial_rate_per_aana') or 0)
    avg_rate  = (govt_rate + comm_rate) / 2
    avg_fmv_rate = _money(avg_rate) if avg_rate else '—'
    # Decimal Aana from stored field, or computed from traditional R-A-P-D
    # Nepali: 1 Ropani = 16 Aana, 1 Aana = 4 Paisa, 1 Paisa = 4 Daam
    def _trad_to_decimal_aana(trad_str):
        try:
            parts = trad_str.split('-')
            if len(parts) == 4:
                r, a, p, d = float(parts[0]), float(parts[1]), float(parts[2]), float(parts[3])
                return round(r * 16 + a + p / 4 + d / 16, 4)
        except Exception:
            pass
        return None

    if prop.get('land_area_aana_decimal'):
        land_aana_dec = _to_float(prop['land_area_aana_decimal'])
    else:
        land_aana_dec = _trad_to_decimal_aana(prop.get('land_area_lorc_trad') or '') or 0.0
    land_aana = str(round(land_aana_dec, 2)) if land_aana_dec else '—'

    # Compute land values: rate per Aana × decimal Aana area
    lorc_area_num = _to_float(prop.get('land_area_lorc') or 0)
    aana_for_calc = land_aana_dec if land_aana_dec else (lorc_area_num / 31.8 if lorc_area_num else 0)
    avg_fmv_rate_val = (govt_rate + comm_rate) / 2

    # Raw computed values (used for 50% breakdown and average)
    govt_land_raw   = govt_rate * aana_for_calc if (aana_for_calc and govt_rate) else 0
    market_land_raw = comm_rate * aana_for_calc if (aana_for_calc and comm_rate) else 0
    govt_land_value   = _money(govt_land_raw) if govt_land_raw else '—'
    market_land_value = _money(market_land_raw) if market_land_raw else '—'

    # 50% Government and 50% Market (for FMV breakdown paragraphs)
    govt_50_raw  = govt_land_raw * 0.5
    market_50_raw = market_land_raw * 0.5
    govt_50   = _money(govt_50_raw) if govt_50_raw else '—'
    market_50 = _money(market_50_raw) if market_50_raw else '—'
    # Average rate per Aana × area = total average land value
    avg_fmv_amount = _money(avg_fmv_rate_val * aana_for_calc) if aana_for_calc and avg_fmv_rate_val else '—'

    # Market values for summary tables; prefer stored, fall back to computed
    mv_land_raw = _to_float(prop.get('market_value_land') or 0)
    mv_bldg_raw = _to_float(prop.get('market_value_building') or 0)
    mv_land     = _money(prop.get('market_value_land')) if prop.get('market_value_land') else market_land_value
    mv_building = _money(prop.get('market_value_building')) if prop.get('market_value_building') else '—'
    mv_total    = _money(mv_land_raw + mv_bldg_raw) if (mv_land_raw or mv_bldg_raw) else (market_land_value if market_land_value != '—' else '—')

    # ── Extra cover-page variables ─────────────────────────────────────────
    def _vdc_ward(vdc, ward):
        if vdc and ward:
            return f'{vdc}-{ward}'
        return vdc or ward or '—'

    def _dist_country(district):
        return f'{district} District, Nepal' if district else '—'

    # Client
    client_vdc_ward = _vdc_ward(c_vdc, c_ward)

    # Owners (individual fields for cover page)
    o1 = owners[0] if owners else {}
    o2 = owners[1] if len(owners) > 1 else {}
    owner1_vdc_ward        = _vdc_ward(_v(o1.get('vdc_municipality'), ''), _v(o1.get('ward_no'), ''))
    owner1_district_country= _dist_country(_v(o1.get('district'), ''))
    owner1_phone           = _v(o1.get('phone') or o1.get('mobile'), '—') if o1 else '—'
    owner2_vdc_ward        = _vdc_ward(_v(o2.get('vdc_municipality'), ''), _v(o2.get('ward_no'), ''))
    owner2_district_country= _dist_country(_v(o2.get('district'), ''))
    owner2_phone           = _v(o2.get('phone') or o2.get('mobile'), '—') if o2 else '—'

    # Sabik display: "Jorpati VDC Sabik"
    sabik_display = f'{sabik_vdc} Sabik' if sabik_vdc else (sabik_str or '—')

    # Firm address split for cover (street / city+country)
    _fparts = [p.strip() for p in firm_address.split(',') if p.strip()]
    if len(_fparts) >= 3:
        firm_street       = ', '.join(_fparts[:-2])
        firm_city_country = ', '.join(_fparts[-2:])
    else:
        firm_street       = firm_address
        firm_city_country = 'Kathmandu, Nepal'

    ctx = {
        # Firm
        'firm_name':         firm_name,
        'firm_address':      firm_address,
        'firm_street':       firm_street,
        'firm_city_country': firm_city_country,
        'firm_phone':        firm_phone,
        'firm_email':        firm_email,

        # Report metadata
        'ref_no':       _v(val.get('ref_no'), '—'),
        'fiscal_year':  _v(val.get('fiscal_year'), '—'),
        'cert_date':    _v(val.get('certification_date'), '—'),
        'inspection_date': _v(val.get('inspection_date'), '—'),
        'valuation_date':  _v(val.get('certification_date'), '—'),

        # Bank
        'bank_name':    bank_name,
        'bank_branch':  _v(val.get('bank_branch'), '—'),
        'bank_address': _v(val.get('bank_address'), '—'),

        # Client (first client)
        'client_entity_type_label':    _client_entity_label(c) if c else '—',
        'client_entity_type':          _v(c.get('entity_type'), 'individual'),
        'client_gender':               _v(c.get('gender'), '—'),
        'client_full_name':            client_name,
        'client_address_line':         client_address_line,
        'client_address_full':         client_address_full,
        'client_district_country':     client_district_country,
        'client_phone':                _v(c.get('phone') or c.get('mobile'), '—'),
        'client_ward_no':              _v(c.get('ward_no'), '—'),
        'client_vdc_municipality':     _v(c.get('vdc_municipality'), '—'),
        'client_district':             _v(c.get('district'), '—'),
        'client_citizenship_no':       _v(c.get('citizenship_no'), '—'),
        'client_citizenship_issued_date':   _v(c.get('citizenship_issued_date'), '—'),
        'client_citizenship_issued_office': _v(c.get('citizenship_issued_office'), '—'),
        'client_father_name':          _v(c.get('father_name'), '—'),
        'client_grandfather_name':     _v(c.get('grandfather_name') or c.get('husband_name'), '—'),
        'client_husband_name':         _v(c.get('husband_name'), '—'),
        'client_pan_no':               _v(c.get('pan_no'), '—'),
        'client_detail_block':         _client_detail_block(c) if c else '—',

        # Owners
        'owners_plain':        _owners_plain(owners),
        'owners_block':        _owners_block(owners),
        'owners_detail_block': _owners_detail_block(owners),
        'owner1_name':         _owner_full_name(owners[0]) if owners else '—',
        'owner2_name':         _owner_full_name(owners[1]) if len(owners) > 1 else '—',

        # Property
        'plot_no':               _v(prop.get('plot_no'), '—'),
        'property_location_full': property_location_full,
        'property_address':      _v(prop.get('address'), '—'),
        'property_ward_no':      p_ward or '—',
        'vdc_municipality':      p_vdc or '—',
        'district':              p_dist or '—',
        'sabik_vdc':             sabik_vdc or '—',
        'sabik_ward_no':         sabik_ward or '—',
        'sabik':                 sabik_str,
        'property_mortgaged':    _v(prop.get('property_mortgaged'), 'Land'),
        'gps_coordinates':       _v(prop.get('gps_coordinates'), '—'),
        'sheet_no':              _v(prop.get('sheet_no'), '—'),
        'location_type':         _v(prop.get('property_type'), 'Residential'),
        'land_shape':            _v(prop.get('land_shape'), '—'),
        'land_level':            _v(prop.get('land_level'), '—'),
        'nature_of_soil':        _v(prop.get('nature_of_soil'), '—'),
        'construction_on_land':  _v(prop.get('construction_on_land'), '—'),
        'positive_features':     _v(prop.get('positive_features'), '—'),
        'negative_features':     _v(prop.get('negative_features'), '—'),
        'nearest_market':        _v(prop.get('nearest_market'), '—'),
        'distance_from_road':    _v(prop.get('nearest_landmark') or prop.get('public_transport_distance'), '—'),
        'frontage':              _v(prop.get('frontage'), '—'),

        # Land area
        'land_area_lorc':       _v(prop.get('land_area_lorc'), '—'),
        'land_area_measured':   _v(prop.get('land_area_measured'), '—'),
        'land_area_deducted':   _v(prop.get('land_area_deducted') or prop.get('land_area_measured'), '—'),
        'land_area_considered': _v(prop.get('considered_area') or prop.get('land_area_deducted') or prop.get('land_area_measured'), '—'),
        'land_area_aana':      land_aana,

        # Traditional R-A-P-D
        'lorc_r': lorc_parts[0], 'lorc_a': lorc_parts[1],
        'lorc_p': lorc_parts[2], 'lorc_d': lorc_parts[3],
        'meas_r': meas_parts[0], 'meas_a': meas_parts[1],
        'meas_p': meas_parts[2], 'meas_d': meas_parts[3],
        'deduct_r': deduct_parts[0], 'deduct_a': deduct_parts[1],
        'deduct_p': deduct_parts[2], 'deduct_d': deduct_parts[3],
        'consid_r': consid_parts[0], 'consid_a': consid_parts[1],
        'consid_p': consid_parts[2], 'consid_d': consid_parts[3],

        # Boundaries
        'north_boundary': _v(prop.get('north_boundary'), '—'),
        'south_boundary': _v(prop.get('south_boundary'), '—'),
        'east_boundary':  _v(prop.get('east_boundary'), '—'),
        'west_boundary':  _v(prop.get('west_boundary'), '—'),
        'boundary_certified_by':  _v(prop.get('legal_reference_no'), '—'),
        'boundary_cert_date':     _v(prop.get('lorc_registration_date'), '—'),

        # Infrastructure
        'motorable_access': _yn(prop.get('motorable_access')),
        'water_supply':     _yn(prop.get('water_supply')),
        'electricity_line': _yn(prop.get('electricity_line')),
        'telephone':        _yn(prop.get('telephone')),
        'tv_cable':         _yn(prop.get('tv_cable')),
        'sewerage':         _yn(prop.get('sewerage')),

        # Influence flags
        **{v: _yn(prop.get(k)) for k, v in bool_fields.items()},

        # Road & access detail
        'road_access':           _road_access(prop),
        'road_access_blueprint': _v(prop.get('road_access_blueprint'), '—'),
        'road_type':             _v(prop.get('road_type'), '—'),
        'road_width':            _v(prop.get('road_width'), '—'),
        'road_side':             _v(prop.get('road_side'), '—'),
        'land_topography':       _v(prop.get('land_topography'), '—'),
        'facing':                _v(prop.get('facing'), '—'),
        'public_transport_distance': _v(prop.get('public_transport_distance'), '—'),
        'tole':                  _v(prop.get('tole'), '—'),
        'accessibility_text':    accessibility_text,
        'area_change':           area_change,
        'land_area_lorc_sqft':   str(lorc_sqft) if lorc_sqft else '—',
        'land_area_meas_sqft':   lm_total_sqft if lm_total_sqft != '—' else (str(meas_sqft) if meas_sqft else '—'),
        'land_area_ded_sqft':    ded_total_sqft if ded_total_sqft != '—' else (str(ded_sqft) if ded_sqft else '—'),
        'land_area_consid_sqft': str(consid_sqft) if consid_sqft else '—',

        # Road type checkboxes (☑ = selected, ☐ = not selected)
        **{f'road_cb_{k}': ('☑' if _v(prop.get('road_type'), '').lower() == label.lower() else '☐')
           for k, label in [
               ('pitched',       'Pitched Road'),
               ('concrete_slab', 'Concrete Slab Road'),
               ('gravelled',     'Gravelled Road'),
               ('earthen',       'Earthen Road'),
               ('block_paved',   'Block Paved Road'),
               ('stone_paved',   'Stone Paved Road'),
               ('foot_trail',    'Foot Trail'),
           ]},

        # Rates
        'commercial_rate': _money(prop.get('commercial_rate_per_aana')),
        'government_rate': _money(prop.get('government_rate_per_aana')),
        'avg_fmv_rate':    avg_fmv_rate,
        'avg_fmv_amount':  avg_fmv_amount,
        'govt_land_value': govt_land_value,
        'govt_50':         govt_50,
        'market_50':       market_50,
        'market_land_value':     mv_land,
        'market_building_value': mv_building,
        'market_value_total':    mv_total,
        'govt_value_remarks':    _v(prop.get('govt_value_remarks'), ''),
        'market_value_remarks':  _v(prop.get('market_value_remarks'), 'As from Local People and the judgement of the valuator'),

        # Valuation
        'fair_market_value':          fmv_total,
        'fair_market_value_land':     fmv_land,
        'fair_market_value_building': fmv_building if fmv_building != '—' else '-',
        'distress_value':             dv_total,
        'dv_land':    _money(prop.get('distress_value_land')),
        'dv_building': _money(prop.get('distress_value_building')),
        'valuation_in_words': _v(prop.get('valuation_in_words'), '—'),

        # Legal — basic
        'ownership_type':          _v(prop.get('ownership_type'), '—'),
        'hold_type':               _v(prop.get('hold_type'), '—'),
        'mode_of_acquisition':     _v(prop.get('mode_of_acquisition'), '—'),
        'lorc_registration_date':  _v(prop.get('lorc_registration_date'), '—'),
        'land_revenue_payment_date': _v(prop.get('land_revenue_payment_date'), '—'),
        # Legal checklist booleans
        'land_revenue_paid':       _yn(prop.get('land_revenue_paid', True)),
        'sale_gift_elapsed':       _yn(prop.get('sale_gift_elapsed', True)),
        'maps_plots_indicated':    _yn(prop.get('maps_plots_indicated', True)),
        'maps_access_marked':      _yn(prop.get('maps_access_marked', True)),
        'maps_shape_tallies':      _yn(prop.get('maps_shape_tallies', True)),
        'boundary_cert_available': _yn(prop.get('boundary_cert_available', True)),
        'boundary_cert_date':      _v(prop.get('boundary_cert_date'), '—'),
        'free_access_available':   _yn(prop.get('free_access_available', True)),
        'acquisition_notice':      _yn(prop.get('acquisition_notice', False)),
        'boundary_clearly_defined': _yn(prop.get('boundary_clearly_defined', True)),
        # Legal comments
        'ownership_comments':          _v(prop.get('ownership_comments'), 'None'),
        'land_revenue_comments':       _v(prop.get('land_revenue_comments'), 'None'),
        'land_registration_comments':  _v(prop.get('land_registration_comments'), 'None'),
        'maps_comments':               _v(prop.get('maps_comments'), 'None'),
        'area_change_comments':        _v(prop.get('area_change_comments'), 'None'),
        'boundary_cert_comments':      _v(prop.get('boundary_cert_comments'), 'None'),
        'general_legal_comments':      _v(prop.get('general_legal_comments'), 'None'),

        # Importance
        'location_importance_text': importance_text,
        'merit_1': merits[0], 'merit_2': merits[1], 'merit_3': merits[2],
        'merit_4': merits[3], 'merit_5': merits[4],
        'property_description': _v(prop.get('notes'), '—'),
        'summary_remarks':      summary_remarks,

        # Commercial importance — prefer explicit user input, fall back to computed
        'proximity_civic_amenities': _v(prop.get('nearest_market'), 'The property is near civic amenities.'),
        'surface_transportation':    'All types of surface transportation can access the property.',
        'land_use':                  f'The land is used for {_v(prop.get("property_type","residential"))} purpose.',
        'value_increase_features':   _v(prop.get('positive_features') or prop.get('notes'), '—'),
        'value_decrease_features':   _v(prop.get('negative_features'), (
            'The Property has no such features.' if not any([
                prop.get('near_army_barracks'), prop.get('near_dumping_site'),
                prop.get('near_hazardous_factory'), prop.get('water_logging')]) else '—')),
        'other_facilities':          'Most kinds of civic amenities are available.',
        'commercial_remarks':        _v(prop.get('notes'), 'N/A.'),
        # Dynamic triangle lists for template loops ({%tr for t in lm_triangles %} etc.)
        'lm_triangles':   [{'label': chr(65 + i), 'side_a': _v(t.get('side_a'),'—'), 'side_b': _v(t.get('side_b'),'—'),
                            'side_c': _v(t.get('side_c'),'—'), 'semi_perimeter': _v(t.get('semi_perimeter'),'—'),
                            'area_sqft': _v(t.get('area_sqft'),'—'), 'aana': _v(t.get('aana'),'—')}
                           for i, t in enumerate(lm_triangles)],
        'ded_triangles':  [{'label': chr(65 + i), 'side_a': _v(t.get('side_a'),'—'), 'side_b': _v(t.get('side_b'),'—'),
                            'side_c': _v(t.get('side_c'),'—'), 'semi_perimeter': _v(t.get('semi_perimeter'),'—'),
                            'area_sqft': _v(t.get('area_sqft'),'—'), 'aana': _v(t.get('aana'),'—')}
                           for i, t in enumerate(ded_triangles)],
        'lorc_triangles': [{'label': chr(65 + i), 'side_a': _v(t.get('side_a'),'—'), 'side_b': _v(t.get('side_b'),'—'),
                            'side_c': _v(t.get('side_c'),'—'), 'semi_perimeter': _v(t.get('semi_perimeter'),'—'),
                            'area_sqft': _v(t.get('area_sqft'),'—'), 'aana': _v(t.get('aana'),'—')}
                           for i, t in enumerate(lorc_triangles)],

        # Flat vars for static template rows (lm_tri_a/b, ded_tri_a/b)
        **{k: v for prefix, tris in [('lm', lm_triangles), ('ded', ded_triangles)]
           for label, idx in [('a', 0), ('b', 1)]
           for t in [tris[idx] if idx < len(tris) else {}]
           for k, v in {
               f'{prefix}_tri_{label}_a':    _v(t.get('side_a'), '—'),
               f'{prefix}_tri_{label}_b':    _v(t.get('side_b'), '—'),
               f'{prefix}_tri_{label}_c':    _v(t.get('side_c'), '—'),
               f'{prefix}_tri_{label}_s':    _v(t.get('semi_perimeter'), '—'),
               f'{prefix}_tri_{label}_sqft': _v(t.get('area_sqft'), '—'),
               f'{prefix}_tri_{label}_aana': _v(t.get('aana'), '—'),
           }.items()},

        # Triangle totals (Section 13 summary rows)
        'lm_total_sqft':  lm_total_sqft,
        'lm_total_sqm':   lm_total_sqm,
        'ded_total_sqft':    ded_total_sqft,
        'ded_total_sqm':     ded_total_sqm,
        'ded_deduct_pct':    str(int(road_ded_pct)) if road_ded_pct == int(road_ded_pct) else str(road_ded_pct),
        'ded_deduct_sqft':   str(road_ded_amt),
        'ded_after_sqft':    ded_after_sqft,
        'ded_after_sqm':     ded_after_sqm,

        # Lalpurja area totals (Section 13C)
        'land_area_lorc_sqm':  lorc_sqm,
        'land_area_lorc_aana': lorc_aana,

        # Certifier
        'certifier_name':   _v(val.get('certifier_name'), '—'),
        'certifier_phone':  _v(val.get('certifier_phone') or val.get('firm_phone'), firm_phone),
        'nec_no':           _v(val.get('nec_no'), '—'),
        'nec_class':        _v(val.get('nec_class'), 'A'),
        'nec_type':         _v(val.get('nec_type'), 'Civil'),
        'site_visited_by':  _v(val.get('site_visited_by'), _v(val.get('certifier_name'), '—')),
        'site_visitor_phone': _v(val.get('site_visitor_phone'), firm_phone),

        # Remarks
        'remarks': _v(val.get('remarks') or prop.get('notes'), '—'),

        # ── Letterhead / property schedule extras ─────────────────────────
        # FMV formatted without trailing .00 for inline table cells
        'fair_market_value_raw': (lambda v: (
            f'{int(float(v)):,}' if v and str(v).replace('.','',1).isdigit() else _money(v)
        ))(prop.get('fair_market_value_total')),
        # Building section of the property table
        'building_status':     (
            prop.get('property_status')
            if prop.get('property_status') and prop.get('property_status') not in ('active', 'Active', '')
            else 'At Present'
        ),
        'building_area_sqft':  _v(prop.get('built_area') or prop.get('total_area'), '-'),
        'building_fmv':        _money(prop.get('fair_market_value_building')) if prop.get('fair_market_value_building') else '-',
        'building_length':     '-',
        'building_breadth':    '-',
        'building_height':     '-',

        # ── Cover-page extras ──────────────────────────────────────────────
        'client_vdc_ward':         client_vdc_ward,
        'owner1_vdc_ward':         owner1_vdc_ward,
        'owner1_district_country': owner1_district_country,
        'owner1_phone':            owner1_phone,
        'owner2_vdc_ward':         owner2_vdc_ward,
        'owner2_district_country': owner2_district_country,
        'owner2_phone':            owner2_phone,
        'sabik_display':           sabik_display,
    }
    return ctx


_LOGO_NAMES = ('firm_logo.png', 'firm_logo.jpg', 'firm_logo.jpeg')

def _find_logo():
    for name in _LOGO_NAMES:
        path = os.path.join(TEMPLATES_DIR, name)
        if os.path.exists(path):
            return path
    return None

def _all_paragraphs(doc):
    """Yield every paragraph in the document, including those in table cells."""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def _inject_logo_post(docx_bytes):
    """
    Post-render step: find the first EMPTY paragraph that follows the
    'Prepared by' heading and place the firm logo there (centred, 2 in wide).
    Searches both top-level paragraphs and table cells.
    """
    logo_path = _find_logo()
    if not logo_path:
        return docx_bytes
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import copy

        doc = Document(io.BytesIO(docx_bytes))

        # Find the 'Prepared by' paragraph, then take the next empty one
        all_paras = list(_all_paragraphs(doc))
        target = None
        found_prep = False
        for para in all_paras:
            txt = para.text.strip().lower()
            if not found_prep:
                if 'prepared by' in txt:
                    found_prep = True
            else:
                # First empty (or near-empty) paragraph after "Prepared by"
                if txt == '':
                    target = para
                    break

        if target is None and found_prep:
            # Fall back: use the paragraph right after "Prepared by"
            for i, para in enumerate(all_paras):
                if 'prepared by' in para.text.strip().lower():
                    if i + 1 < len(all_paras):
                        target = all_paras[i + 1]
                    break

        if target is not None:
            target.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Clear existing runs
            for run in list(target.runs):
                run._element.getparent().remove(run._element)
            run = target.add_run()
            run.add_picture(logo_path, width=Inches(2.0))

        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()
    except Exception:
        return docx_bytes


class DocumentService:

    def generate(self, doc_type, hierarchy, valuation):
        template_file = TEMPLATE_FILES.get(doc_type)
        if not template_file:
            raise ValueError(f'Unknown document type: {doc_type}')
        template_path = os.path.join(TEMPLATES_DIR, template_file)
        if not os.path.exists(template_path):
            raise FileNotFoundError(f'Template not found: {template_path}')
        tpl = DocxTemplate(template_path)
        context = build_context(hierarchy, valuation)
        tpl.render(context)
        buf = io.BytesIO()
        tpl.save(buf)
        raw = buf.getvalue()

        # Inject logo into cover / proposal
        if doc_type in ('cover', 'proposal'):
            raw = _inject_logo_post(raw)

        return raw
