"""
Build .docx template files. Run: python build_templates.py
Uses docxtpl Jinja2 syntax: {{ var }}, {% for %}, etc.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DIR = os.path.dirname(os.path.abspath(__file__))

# ── Helpers ────────────────────────────────────────────────────────────────────

def _shd(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), hex_color)
    tcPr.append(s)

def _border(cell, sides=('top','bottom','left','right'), size=6, color='000000'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def p(doc, text='', bold=False, size=10, align=None, space_before=0, space_after=0, italic=False, color=None):
    para = doc.add_paragraph()
    if align:
        para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if text:
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
    return para

def p2(doc, parts, align=None, space_before=0, space_after=2):
    """parts = [(text, bold, size, italic, color), ...]"""
    para = doc.add_paragraph()
    if align:
        para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    for part in parts:
        text = part[0]
        bold = part[1] if len(part) > 1 else False
        size = part[2] if len(part) > 2 else 10
        italic = part[3] if len(part) > 3 else False
        color = part[4] if len(part) > 4 else None
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
    return para

def hr(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '1F3A2E')
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_after = Pt(4)

def page_break(doc):
    doc.add_page_break()

def set_col_widths(table, widths):
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)

def tbl_row(table, cells, bold_first=False, shd_color=None):
    row = table.add_row()
    for i, (cell, txt) in enumerate(zip(row.cells, cells)):
        cell.text = txt
        cell.paragraphs[0].runs[0].font.size = Pt(9.5) if cell.paragraphs[0].runs else None
        if shd_color:
            _shd(cell, shd_color)
        if bold_first and i == 0:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
    return row

def header_row(table, cols, bg='1F3A2E', txt_color='FFFFFF'):
    row = table.rows[0]
    for cell, txt in zip(row.cells, cols):
        cell.text = txt
        if cell.paragraphs[0].runs:
            r = cell.paragraphs[0].runs[0]
            r.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(*bytes.fromhex(txt_color))
        _shd(cell, bg)


# ── Firm header block (reused in all docs) ────────────────────────────────────

def add_firm_header(doc):
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hdr.add_run('{{ firm_name }}')
    r.bold = True; r.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('{{ firm_address }}')
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(80,80,80)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run('Tel: {{ firm_phone }}   |   Email: {{ firm_email }}')
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(80,80,80)
    hr(doc)


# ── 1. COVER TEMPLATE ─────────────────────────────────────────────────────────

def build_cover():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1.0); sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.2); sec.right_margin = Inches(1.2)

    add_firm_header(doc)

    p(doc, '', space_before=20)

    p(doc, 'Property Valuation Report', bold=True, size=18,
      align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10)
    p(doc, '({{ property_mortgaged }})', size=12,
      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    hr(doc)

    # Client block
    p(doc, 'Client', bold=True, size=11, space_before=8)
    for field in ['{{ client_full_name }}','{{ client_address_line }}',
                  '{{ client_district_country }}','Contact No: {{ client_phone }}']:
        p(doc, field, size=10, space_after=1)

    p(doc, '', space_before=10)
    p(doc, 'Properties to be Valued', bold=True, size=11)

    p2(doc, [('Plot No.: ', True, 10), ('{{ plot_no }}', False, 10)])
    p(doc, '{{ property_location_full }}', size=10, space_after=1)
    p(doc, '{{ sabik }}', size=10, italic=True)

    # Owners block (text merge – multiple owners rendered in template)
    p(doc, '', space_before=8)
    p(doc, '{{ owners_block }}', size=10)

    p(doc, '', space_before=10)
    hr(doc)

    p(doc, 'Submitted to', bold=True, size=11, space_before=8)
    p(doc, '{{ bank_name }}', bold=True, size=10, space_after=1)
    p(doc, '{{ bank_branch }}', size=10, space_after=1)

    p(doc, '', space_before=10)
    p2(doc, [('Fair Market Value of the Land & Building: NRs. ', False, 11),
              ('{{ fair_market_value }}', True, 11)])

    p(doc, '', space_before=14)
    hr(doc)

    p(doc, 'Prepared by', bold=True, size=11, space_before=8)
    for field in ['{{ firm_name }}','{{ firm_address }}','Tel: {{ firm_phone }}',
                  'Email: {{ firm_email }}']:
        p(doc, field, size=10, space_after=1)

    p(doc, '', space_before=14)
    p(doc, '{{ cert_date }}', size=10,
      align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(os.path.join(DIR, 'cover_template.docx'))
    print('  ✓ cover_template.docx')


# ── 2. LETTERHEAD TEMPLATE ────────────────────────────────────────────────────

def build_letterhead():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(1.0); sec.right_margin = Inches(1.0)

    add_firm_header(doc)

    # Ref + Date
    ref_p = doc.add_paragraph()
    ref_p.paragraph_format.space_after = Pt(0)
    ref_p.add_run('Ref No.: ').bold = True
    ref_p.runs[-1].font.size = Pt(10)
    ref_p.add_run('{{ ref_no }}').font.size = Pt(10)

    date_p = doc.add_paragraph()
    date_p.paragraph_format.space_after = Pt(8)
    r = date_p.add_run('Date: {{ cert_date }}')
    r.font.size = Pt(10)

    p(doc, 'To', bold=True, size=10, space_after=1)
    for field in ['The Respected Manager,','{{ bank_name }},',
                  '{{ bank_branch }},','{{ bank_address }}']:
        p(doc, field, size=10, space_after=1)

    p(doc, '', space_before=4)
    p(doc, 'Dear Sir/Madam,', size=10)
    p(doc, '', space_before=4)
    p(doc, 'Valuation Certificate', bold=True, size=12,
      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    p(doc, '            We are submitting herewith the valuation report of the property '
           'which are intended to be mortgaged to the bank by:', size=10, space_after=4)

    # Client info
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = 'Table Grid'
    set_col_widths(tbl, [1.2, 5.0])
    tbl.rows[0].cells[0].text = 'Client'
    tbl.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    tbl.rows[0].cells[1].text = '{{ client_full_name }}'
    tbl.rows[1].cells[0].text = 'Address:'
    tbl.rows[1].cells[0].paragraphs[0].runs[0].bold = True
    tbl.rows[1].cells[1].text = '{{ client_address_full }}'
    for row in tbl.rows:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10) if cell.paragraphs[0].runs else None

    p(doc, '', space_before=6)
    p(doc, 'We hereby declare and certify that:', bold=True, size=10, space_after=3)

    certify_points = [
        '1.  We have physically inspected, verified & measured the properties on {{ inspection_date }}.',
        '2.  We have no direct or indirect interests in the said company or properties.',
        '3.  The information furnished are true & correct to the best made available to us by the client.',
        '4.  The value of the properties as detailed in this report has been carried out by us in strict '
             'compliance with the "Valuation Guidelines -2068" of {{ bank_name }}.',
        '5.  The said property is acceptable security to the Bank in all respects.',
    ]
    for pt in certify_points:
        p(doc, pt, size=10, space_after=2)

    p(doc, '', space_before=4)

    # Summary table
    stbl = doc.add_table(rows=1, cols=8)
    stbl.style = 'Table Grid'
    header_row(stbl, ['S.N.','Owner of the property','Location of Property',
                       'Plot No.','Area as per\nLalpurja (Sq.m)',
                       'Area as per\nMeasurement (Sq.m)',
                       'Fair Market\nValue (NRs.)',
                       'Building\nStatus'])
    set_col_widths(stbl, [0.35, 1.2, 1.5, 0.55, 0.85, 0.85, 1.0, 0.7])

    data_row = stbl.add_row()
    vals = ['1', '{{ owners_plain }}', '{{ property_location_full }}',
            '{{ plot_no }}', '{{ land_area_lorc }}',
            '{{ land_area_measured }}', '{{ fair_market_value }}', 'At Present']
    for cell, v in zip(data_row.cells, vals):
        cell.text = v
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    area_row = stbl.add_row()
    area_row.cells[4].text = 'Area Considered'
    if area_row.cells[4].paragraphs[0].runs:
        area_row.cells[4].paragraphs[0].runs[0].font.size = Pt(9)
        area_row.cells[4].paragraphs[0].runs[0].italic = True
    area_row.cells[5].text = '{{ land_area_considered }}'
    if area_row.cells[5].paragraphs[0].runs:
        area_row.cells[5].paragraphs[0].runs[0].font.size = Pt(9)

    sheet_row = stbl.add_row()
    sheet_row.cells[1].text = 'Sheet No: {{ sheet_no }}'
    if sheet_row.cells[1].paragraphs[0].runs:
        sheet_row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

    p(doc, '', space_before=4)

    p2(doc, [('Total Fair Market Value of Land & Building Project', True, 10),
              (' (NRs.): ', False, 10),
              ('{{ fair_market_value }}', True, 10)])

    p2(doc, [('FAIR MARKET VALUE OF THE PROPERTY (NRs.): ', True, 10),
              ('{{ fair_market_value }}', True, 10)])

    p2(doc, [('In Words: ', False, 10),
              ('{{ valuation_in_words }}', False, 10, True)])

    p(doc, 'All necessary calculations and documents are attached herewith for ready reference and record.',
      size=10, space_before=4)

    p(doc, '', space_before=10)
    p(doc, 'Yours faithfully,', size=10, space_after=18)
    p(doc, '{{ certifier_name }}', bold=True, size=10, space_after=1)
    p(doc, 'NEC. No.: {{ nec_no }} "Civil" "{{ nec_class }}" Class', size=10, space_after=1)
    p(doc, 'For {{ firm_name }}', size=10)

    doc.save(os.path.join(DIR, 'letterhead_template.docx'))
    print('  ✓ letterhead_template.docx')


# ── 3. PROPOSAL (FULL REPORT) TEMPLATE ───────────────────────────────────────

def build_proposal():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1.0); sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.2); sec.right_margin = Inches(1.0)

    # ── PAGE 1: Cover ──────────────────────────────────────────────────────────
    add_firm_header(doc)
    p(doc, '', space_before=16)
    p(doc, 'Property Valuation Report', bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '({{ property_mortgaged }})', size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    hr(doc)

    p(doc, 'Client', bold=True, size=11, space_before=8)
    for f in ['{{ client_full_name }}','{{ client_address_line }}',
              '{{ client_district_country }}','Contact No: {{ client_phone }}']:
        p(doc, f, size=10, space_after=1)

    p(doc, '', space_before=8)
    p(doc, 'Properties to be Valued', bold=True, size=11)
    p2(doc, [('Plot No.: ', True, 10), ('{{ plot_no }}', False, 10)])
    p(doc, '{{ property_location_full }}', size=10)
    p(doc, '{{ sabik }}', size=10, italic=True)

    p(doc, '', space_before=8)
    p(doc, '{{ owners_block }}', size=10)

    p(doc, '', space_before=10); hr(doc)
    p(doc, 'Submitted to', bold=True, size=11, space_before=8)
    p(doc, '{{ bank_name }}', bold=True, size=10, space_after=1)
    p(doc, '{{ bank_branch }}', size=10)

    p(doc, '', space_before=10)
    p2(doc, [('Fair Market Value of the Land & Building: NRs. ', False, 11),
              ('{{ fair_market_value }}', True, 11)])

    p(doc, '', space_before=12); hr(doc)
    p(doc, 'Prepared by', bold=True, size=11, space_before=8)
    for f in ['{{ firm_name }}','{{ firm_address }}','Tel: {{ firm_phone }}','Email: {{ firm_email }}']:
        p(doc, f, size=10, space_after=1)
    p(doc, '', space_before=10)
    p(doc, '{{ cert_date }}', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── PAGE 2: Letterhead ─────────────────────────────────────────────────────
    page_break(doc)
    add_firm_header(doc)

    p2(doc, [('Ref No.: ', True, 10), ('{{ ref_no }}', False, 10)])
    p2(doc, [('Date: ', False, 10), ('{{ cert_date }}', False, 10)], space_after=6)

    p(doc, 'To', bold=True, size=10, space_after=1)
    for f in ['The Respected Manager,','{{ bank_name }},','{{ bank_branch }},','{{ bank_address }}']:
        p(doc, f, size=10, space_after=1)

    p(doc, '', space_before=4)
    p(doc, 'Dear Sir/Madam,', size=10)
    p(doc, 'Valuation Certificate', bold=True, size=12,
      align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)

    p(doc, '            We are submitting herewith the valuation report of the property '
           'which are intended to be mortgaged to the bank by:', size=10, space_after=4)

    lt = doc.add_table(rows=2, cols=2); lt.style = 'Table Grid'
    set_col_widths(lt, [1.2, 5.0])
    lt.rows[0].cells[0].text = 'Client'; lt.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    lt.rows[0].cells[1].text = '{{ client_full_name }}'
    lt.rows[1].cells[0].text = 'Address:'; lt.rows[1].cells[0].paragraphs[0].runs[0].bold = True
    lt.rows[1].cells[1].text = '{{ client_address_full }}'
    for row in lt.rows:
        for cell in row.cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)

    p(doc, '', space_before=6)
    p(doc, 'We hereby declare and certify that:', bold=True, size=10, space_after=3)
    for pt in [
        '1.  We have physically inspected, verified & measured the properties on {{ inspection_date }}.',
        '2.  We have no direct or indirect interests in the said company or properties.',
        '3.  The information furnished are true & correct to the best made available to us by the client.',
        '4.  The value of the properties as detailed in this report has been carried out by us in strict compliance with the "Valuation Guidelines -2068" of {{ bank_name }}.',
        '5.  The said property is acceptable security to the Bank in all respects.',
    ]:
        p(doc, pt, size=10, space_after=2)

    p(doc, '', space_before=4)
    st = doc.add_table(rows=1, cols=8); st.style = 'Table Grid'
    header_row(st, ['S.N.','Owner','Location','Plot No.',
                    'Lalpurja\n(Sq.m)','Measured\n(Sq.m)','FMV (NRs.)','Building'])
    set_col_widths(st, [0.35, 1.3, 1.5, 0.55, 0.75, 0.75, 0.9, 0.7])
    dr = st.add_row()
    for cell, v in zip(dr.cells,
            ['1','{{ owners_plain }}','{{ property_location_full }}','{{ plot_no }}',
             '{{ land_area_lorc }}','{{ land_area_measured }}','{{ fair_market_value }}','At Present']):
        cell.text = v
        if cell.paragraphs[0].runs: cell.paragraphs[0].runs[0].font.size = Pt(9)
    ar = st.add_row()
    ar.cells[4].text = 'Area Considered'
    ar.cells[5].text = '{{ land_area_considered }}'
    for cell in ar.cells:
        if cell.paragraphs[0].runs: cell.paragraphs[0].runs[0].font.size = Pt(9)
    sr = st.add_row()
    sr.cells[1].text = 'Sheet No: {{ sheet_no }}'
    if sr.cells[1].paragraphs[0].runs: sr.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

    p(doc, '', space_before=4)
    p2(doc, [('FAIR MARKET VALUE OF THE PROPERTY (NRs.): ', True, 10), ('{{ fair_market_value }}', True, 10)])
    p2(doc, [('In Words: ', False, 10), ('{{ valuation_in_words }}', False, 10, True)])
    p(doc, 'All necessary calculations and documents are attached herewith for ready reference and record.',
      size=10, space_before=4)
    p(doc, '', space_before=10)
    p(doc, 'Yours faithfully,', size=10, space_after=16)
    p(doc, '{{ certifier_name }}', bold=True, size=10, space_after=1)
    p(doc, 'NEC. No.: {{ nec_no }} "Civil" "{{ nec_class }}" Class', size=10, space_after=1)
    p(doc, 'For {{ firm_name }}', size=10)

    # ── PAGE 3: TOC ────────────────────────────────────────────────────────────
    page_break(doc)
    p(doc, 'Synopsis of Property Valuation Report', bold=True, size=13, space_after=6)
    toc_items = [
        'Letter of Bank & Covering Letter of the Consultant Engineer.',
        '1. General Information',
        '   a. Name, Address and General Information of the Client',
        '   b. Property to be Mortgaged',
        '   c. Name, Address and General Information of the Owner of the Property',
        '   d. Location of the Property',
        '   e. Total Area of the Land',
        '   f. Parameters of the Four Boundaries',
        '   g. Importance of the Location',
        '   h. Description of the Property',
        '2. Legal Aspects of the Property',
        '3. Summary Declaration Sheet',
        '4. Commercial Importance of the Land and Building',
        '5. Summary of Valuation of the Property',
        '6. Valuation of the Property',
        '7. Limiting Conditions',
        '8. Remarks',
        '9. Opinion',
        '10. Declaration',
        '11. Acceptance and Declaration of the Property by the Owner and the Client',
        '12. Certification',
        '13. Site Measurement of the Land (Area Calculation)',
        '14. Methodology of Valuation',
        '15. Terminology Used',
        '16. Property Location Map / Triangulation',
        'Annex A: GPS Information & Photographs of the Property',
        'Annex B: Supporting Legal Documents',
        '   1. Copy of Citizenship Certificate of the Owner/Client',
        '   2. Copy of Title Deed Certificate (Lalpurja)',
        '   3. Land Revenue Payment Receipt',
        '   4. Four Boundary Disclosure',
        '   5. Original Blueprint/Trace Map',
    ]
    for item in toc_items:
        p(doc, item, size=10, space_after=1)

    # ── PAGE 4: Synopsis (Section 1) ───────────────────────────────────────────
    page_break(doc)
    p(doc, '1. General Information', bold=True, size=12, space_after=6)

    # Section a – Client
    p(doc, 'a. Name, Address and General Information of the Client', bold=True, size=10, space_after=4)
    ct = doc.add_table(rows=1, cols=2); ct.style = 'Table Grid'
    set_col_widths(ct, [2.2, 4.0])
    header_row(ct, ['Field', 'Details'])
    client_rows = [
        ('Client', '{{ client_full_name }}'),
        ('Ward No.', '{{ client_ward_no }}'),
        ('Municipality / VDC', '{{ client_vdc_municipality }}'),
        ('District', '{{ client_district }}'),
        ('Citizenship No.', '{{ client_citizenship_no }}'),
        ('Issued Date', '{{ client_citizenship_issued_date }}'),
        ('Issued Office', '{{ client_citizenship_issued_office }}'),
        ("Father's Name", '{{ client_father_name }}'),
        ("Grand Father's / Husband's Name", '{{ client_grandfather_name }}'),
        ('Contact No.', '{{ client_phone }}'),
    ]
    for label, val in client_rows:
        tbl_row(ct, [label, val])

    # Section b – Property to be Mortgaged
    p(doc, '', space_before=8)
    p2(doc, [('b. Property to be Mortgaged: ', True, 10), ('{{ property_mortgaged }}', False, 10)])

    # Section c – Owners
    p(doc, '', space_before=8)
    p(doc, 'c. Name, Address and General Information of the Owner(s)', bold=True, size=10, space_after=4)
    p(doc, '{{ owners_detail_block }}', size=10)

    # ── PAGE 5: Location ───────────────────────────────────────────────────────
    page_break(doc)
    p(doc, 'd. Location of the Property', bold=True, size=12, space_after=6)

    loct = doc.add_table(rows=1, cols=2); loct.style = 'Table Grid'
    set_col_widths(loct, [2.5, 3.7])
    header_row(loct, ['Field', 'Details'])
    loc_rows = [
        ('Plot No.', '{{ plot_no }}'),
        ('Ward No.', '{{ property_ward_no }}'),
        ('VDC / Municipality', '{{ vdc_municipality }}'),
        ('District', '{{ district }}'),
        ('Type of Location', '{{ location_type }}'),
        ('Shape', '{{ land_shape }}'),
        ('Nearest Market', '{{ nearest_market }}'),
        ('Distance from Main Road', '{{ distance_from_road }}'),
        ('High Tension Line', '{{ near_high_tension_line_text }}'),
        ('GPS Co-ordinate', '{{ gps_coordinates }}'),
        ('Sheet No.', '{{ sheet_no }}'),
    ]
    for label, val in loc_rows:
        tbl_row(loct, [label, val])

    p(doc, '', space_before=8)
    p(doc, 'e. Total Area of Land', bold=True, size=10, space_after=4)
    areat = doc.add_table(rows=1, cols=6); areat.style = 'Table Grid'
    header_row(areat, ['', 'Ropani', 'Aana', 'Paisa', 'Daam', 'Sq.m.'])
    set_col_widths(areat, [1.6, 0.7, 0.7, 0.7, 0.7, 0.9])
    area_rows = [
        ('As per Lalpurja', '{{ lorc_r }}', '{{ lorc_a }}', '{{ lorc_p }}', '{{ lorc_d }}', '{{ land_area_lorc }}'),
        ('As per Measurement', '{{ meas_r }}', '{{ meas_a }}', '{{ meas_p }}', '{{ meas_d }}', '{{ land_area_measured }}'),
        ('After Deduction', '{{ deduct_r }}', '{{ deduct_a }}', '{{ deduct_p }}', '{{ deduct_d }}', '{{ land_area_deducted }}'),
        ('Area Considered for Valuation', '{{ consid_r }}', '{{ consid_a }}', '{{ consid_p }}', '{{ consid_d }}', '{{ land_area_considered }}'),
    ]
    for row_data in area_rows:
        tbl_row(areat, list(row_data), bold_first=True)

    p(doc, '', space_before=8)
    p(doc, 'f. Parameters of the Four Boundaries', bold=True, size=10, space_after=4)
    boundary_cert = p(doc, 'Certified By: {{ boundary_certified_by }}', size=9, italic=True, space_after=4)

    bt = doc.add_table(rows=1, cols=5); bt.style = 'Table Grid'
    header_row(bt, ['Plot No.', 'East', 'West', 'North', 'South'])
    tbl_row(bt, ['{{ plot_no }}', '{{ east_boundary }}', '{{ west_boundary }}',
                  '{{ north_boundary }}', '{{ south_boundary }}'])

    # ── PAGE 6: Importance ─────────────────────────────────────────────────────
    page_break(doc)
    p(doc, 'g. Importance of the Location', bold=True, size=12, space_after=6)
    p(doc, '{{ location_importance_text }}', size=10, space_after=6)
    p(doc, 'Following are some of the merits of the location:', size=10, space_before=4)
    for bullet in ['{{ merit_1 }}','{{ merit_2 }}','{{ merit_3 }}','{{ merit_4 }}','{{ merit_5 }}']:
        p2(doc, [('•  ', False, 10), (bullet, False, 10)], space_after=2)

    p(doc, '', space_before=8)
    p(doc, 'h. Description of the Property', bold=True, size=12, space_after=4)
    p(doc, '{{ property_description }}', size=10)

    # ── PAGE 7: Legal Aspects ──────────────────────────────────────────────────
    page_break(doc)
    p(doc, '2. Legal Aspects of the Property', bold=True, size=12, space_after=4)
    p(doc, 'A number of local people and landowners in the vicinity were interviewed and to the '
           'best of our knowledge, no legal dispute has been recorded. Ownership of property has '
           'been verified on the basis of Title Deed Certificate (Lalpurja), interview with Client '
           'and the Local People during the field survey in {{ inspection_date }}.', size=10, space_after=6)

    legal_sections = {
        'Land Ownership Document (Lalpurja)': [
            ('a) Type of Ownership', '{{ ownership_type }}'),
            ('b) Ownership of Land', '{{ hold_type }}'),
            ('c) Name of the Owner', '{{ owners_plain }}'),
            ('d) Comments, if any', 'None'),
        ],
        'Land Revenue (Malpot)': [
            ('a) Current Revenue has been paid', 'Yes'),
            ('b) Date of Payment of Receipt', '{{ land_revenue_payment_date }}'),
            ('c) Comments, if any', 'None'),
        ],
        'Land Registration Paper': [
            ('a) Whether Normal Sale/Gift', '{{ mode_of_acquisition }}'),
            ('b) Date of Registration', '{{ lorc_registration_date }}'),
            ('c) Whether 6 Month 35 days has elapsed', 'Yes'),
            ('d) Comments, if any', 'None'),
        ],
        'Certificate of Parameters of Boundaries': [
            ('a) Parameters of Boundary Available', 'Yes'),
            ('b) Date of Certification', '{{ boundary_cert_date }}'),
            ('c) Comments, if any', 'None'),
        ],
        'General': [
            ('a) Free access to the property is Available', 'Yes'),
            ('b) Land notified for acquisition by Government', 'No'),
            ('c) Boundary clearly defined at site', 'Yes'),
            ('d) Comments, if any', 'None'),
        ],
    }
    for section_title, items in legal_sections.items():
        p(doc, section_title, bold=True, size=10, space_before=6, space_after=3)
        lt2 = doc.add_table(rows=1, cols=2); lt2.style = 'Table Grid'
        header_row(lt2, ['Item', 'Status'])
        set_col_widths(lt2, [4.0, 2.2])
        for label, val in items:
            tbl_row(lt2, [label, val])

    # ── PAGE 8: Summary Declaration ────────────────────────────────────────────
    page_break(doc)
    p(doc, '3. Summary Declaration Sheet', bold=True, size=12, space_after=6)

    infra_checks = [
        ('Motorable Access',               '{{ motorable_access }}'),
        ('Water Supply Line',              '{{ water_supply }}'),
        ('Electricity Line',               '{{ electricity_line }}'),
        ('Telephone Line',                 '{{ telephone }}'),
        ('TV Cable',                       '{{ tv_cable }}'),
        ('Sewerage Pipe Line',             '{{ sewerage }}'),
        ('River/Stream near the property', '{{ near_river_stream_text }}'),
        ('Fuel storage depot',             '{{ near_fuel_depot_text }}'),
        ('Temple/Shrine near property',    '{{ near_temple_text }}'),
        ('Water Logging',                  '{{ water_logging_text }}'),
        ('Cremation area near property',   '{{ near_cremation_area_text }}'),
        ('Army barracks near property',    '{{ near_army_barracks_text }}'),
        ('Monument near property',         '{{ near_monument_text }}'),
        ('Hazardous factory',              '{{ near_hazardous_factory_text }}'),
        ('Dumping site near property',     '{{ near_dumping_site_text }}'),
        ('High-tension line near property','{{ near_high_tension_line_text }}'),
    ]
    inft = doc.add_table(rows=1, cols=2); inft.style = 'Table Grid'
    header_row(inft, ['Item', 'Status'])
    set_col_widths(inft, [3.8, 2.4])
    for label, val in infra_checks:
        tbl_row(inft, [label, val])

    p(doc, '', space_before=6)
    p(doc, '4. Commercial Importance of the Land and Building', bold=True, size=12, space_after=4)
    commercial_items = [
        ('Proximity to Civic Amenities', '{{ proximity_civic_amenities }}'),
        ('Surface Transportation', '{{ surface_transportation }}'),
        ('Use of Land', '{{ land_use }}'),
        ('Features for Increase of Value', '{{ value_increase_features }}'),
        ('Features for Decrease of Value', '{{ value_decrease_features }}'),
        ('Other Facilities Available', '{{ other_facilities }}'),
        ('Remarks', '{{ commercial_remarks }}'),
    ]
    comt = doc.add_table(rows=1, cols=2); comt.style = 'Table Grid'
    header_row(comt, ['Item', 'Details'])
    set_col_widths(comt, [2.5, 3.7])
    for label, val in commercial_items:
        tbl_row(comt, [label, val])

    # ── PAGE 9: Summary of Valuation ───────────────────────────────────────────
    page_break(doc)
    p(doc, '5. Summary of Valuation of the Property', bold=True, size=12, space_after=4)

    p2(doc, [('Owner(s): ', True, 10), ('{{ owners_plain }}', False, 10)])
    p2(doc, [('Date of Survey: ', True, 10), ('{{ inspection_date }}', False, 10)])
    p2(doc, [('Type of Property: ', True, 10), ('{{ property_mortgaged }}', False, 10)])
    p2(doc, [('Plot No.: ', True, 10), ('{{ plot_no }}', False, 10)])
    p2(doc, [('Location: ', True, 10), ('{{ property_location_full }}', False, 10)])
    p2(doc, [('Client: ', True, 10), ('{{ client_full_name }}', False, 10)])

    p(doc, '', space_before=6)
    p(doc, 'Area of the Land', bold=True, size=10, space_after=4)
    sa = doc.add_table(rows=1, cols=6); sa.style = 'Table Grid'
    header_row(sa, ['', 'R', 'A', 'P', 'D', 'Sq.m.'])
    set_col_widths(sa, [1.8, 0.5, 0.5, 0.5, 0.5, 0.95])
    for row_data in [
        ('As per Lalpurja', '{{ lorc_r }}', '{{ lorc_a }}', '{{ lorc_p }}', '{{ lorc_d }}', '{{ land_area_lorc }}'),
        ('As per Measurement', '{{ meas_r }}', '{{ meas_a }}', '{{ meas_p }}', '{{ meas_d }}', '{{ land_area_measured }}'),
        ('After Deduction', '{{ deduct_r }}', '{{ deduct_a }}', '{{ deduct_p }}', '{{ deduct_d }}', '{{ land_area_deducted }}'),
        ('Area Considered', '{{ consid_r }}', '{{ consid_a }}', '{{ consid_p }}', '{{ consid_d }}', '{{ land_area_considered }}'),
    ]:
        tbl_row(sa, list(row_data), bold_first=True)

    p(doc, '', space_before=6)
    p(doc, 'Parameters of Four Boundaries', bold=True, size=10, space_after=4)
    bbt = doc.add_table(rows=1, cols=5); bbt.style = 'Table Grid'
    header_row(bbt, ['Plot', 'East', 'West', 'North', 'South'])
    tbl_row(bbt, ['{{ plot_no }}', '{{ east_boundary }}', '{{ west_boundary }}',
                  '{{ north_boundary }}', '{{ south_boundary }}'])

    p(doc, '', space_before=6)
    p(doc, '{{ summary_remarks }}', size=10)

    p(doc, '', space_before=6)
    p(doc, 'Value of the Property', bold=True, size=10, space_after=4)
    vt = doc.add_table(rows=1, cols=3); vt.style = 'Table Grid'
    header_row(vt, ['S.N.', 'Particulars', 'Fair Market Value (NRs.)'])
    set_col_widths(vt, [0.4, 2.4, 2.0])
    tbl_row(vt, ['1', 'Land', '{{ fair_market_value_land }}'])
    tbl_row(vt, ['2', 'Building', '{{ fair_market_value_building }}'])
    tbl_row(vt, ['', 'Total Amount', '{{ fair_market_value }}'])

    p2(doc, [('On the basis of the details provided by Client and Field verification, '
              'Fair Market Value of property to be mortgaged is valued NRs. ', False, 10),
              ('{{ fair_market_value }}', True, 10),
              ('.', False, 10)], space_before=6)
    p2(doc, [('In Words: ', False, 10), ('{{ valuation_in_words }}', False, 10, True)])

    # ── PAGE 10: Valuation ─────────────────────────────────────────────────────
    page_break(doc)
    p(doc, '6. Valuation of the Property', bold=True, size=12, space_after=4)

    p2(doc, [('Owner(s): ', True, 10), ('{{ owners_plain }}', False, 10)])
    p2(doc, [('Date of Survey: ', True, 10), ('{{ inspection_date }}', False, 10)])
    p2(doc, [('Type of Property: ', True, 10), ('{{ property_mortgaged }}', False, 10)])
    p2(doc, [('Plot No.: ', True, 10), ('{{ plot_no }}', False, 10)])
    p2(doc, [('Location: ', True, 10), ('{{ property_location_full }}', False, 10)])
    p2(doc, [('Client: ', True, 10), ('{{ client_full_name }}', False, 10)])

    p(doc, '', space_before=6)
    p(doc, 'Value of the Land', bold=True, size=10, space_after=4)
    vlandt = doc.add_table(rows=1, cols=5); vlandt.style = 'Table Grid'
    header_row(vlandt, ['S.N.', 'Particulars', 'Area (Aana)', 'Rate per Aana (NRs.)', 'Amount (NRs.)'])
    set_col_widths(vlandt, [0.4, 2.0, 0.9, 1.5, 1.5])
    tbl_row(vlandt, ['1', 'Total Government Value of Land', '{{ land_area_aana }}',
                     '{{ government_rate }}', '{{ govt_land_value }}'])
    tbl_row(vlandt, ['2', 'Total Market Value of Land', '{{ land_area_aana }}',
                     '{{ commercial_rate }}', '{{ market_land_value }}'])
    tbl_row(vlandt, ['', 'Fair Market Value of Land (50% Govt + 50% Market)',
                     '', '', '{{ fair_market_value_land }}'])
    tbl_row(vlandt, ['', 'Average FMV Rate per Aana', '', '{{ avg_fmv_rate }}', ''])

    p(doc, '', space_before=6)
    p(doc, 'VALUATION OF THE PROPERTY', bold=True, size=10, space_after=4)
    vtott = doc.add_table(rows=1, cols=4); vtott.style = 'Table Grid'
    header_row(vtott, ['S.N.', 'Particulars', 'Commercial Value (NRs.)', 'Fair Market Value (NRs.)'])
    set_col_widths(vtott, [0.4, 1.8, 2.0, 2.0])
    tbl_row(vtott, ['1', 'Land', '{{ market_land_value }}', '{{ fair_market_value_land }}'])
    tbl_row(vtott, ['2', 'Building', '{{ fair_market_value_building }}', '{{ fair_market_value_building }}'])
    tbl_row(vtott, ['', 'Total', '{{ market_land_value }}', '{{ fair_market_value }}'])

    p2(doc, [('Fair Market Value of Property (NRs.): ', True, 10), ('{{ fair_market_value }}', True, 10)],
       space_before=6)
    p2(doc, [('In Words: ', False, 10), ('{{ valuation_in_words }}', False, 10, True)])

    # ── PAGE 11: Limiting Conditions ───────────────────────────────────────────
    page_break(doc)
    p(doc, '7. Limiting Conditions', bold=True, size=12, space_after=4)
    for lc in [
        'a)  The opinion of value stated is based on the facts and assumptions identified in this report. '
             'The valuator takes no responsibility for changes in market conditions.',
        'b)  The opinion of value stated is based on the information that has been available to the valuator '
             'at the time; the valuation analysis has been conducted. The value of the property may change '
             'substantially with time and the valuator reserves the right to alter the opinion of value stated '
             'if relevant information is available later.',
    ]:
        p(doc, lc, size=10, space_after=4)

    p(doc, '8. Remarks', bold=True, size=12, space_before=6, space_after=4)
    for remark in [
        'a.  The stated opinions of value are effective as of the date of value based upon information that '
             'was available to the valuer at the time of the valuation analysis was conducted.',
        'b.  To the best of our knowledge, all matters of factual nature discussed in this report are true '
             'and correct. No important factors have been intentionally overlooked or withheld.',
        'c.  It has been fully ascertained that lending of the Bank is fully secured by this mortgaged up '
             'to the Distress Value of the amount certified on the report.',
        'd.  The use of the business property does not indicate any possibility of toxic contamination of '
             'the ground. Our analysis is based upon the assumption that there is no contamination present.',
        'e.  Economical obsolescence of the property has not been considered in this report.',
    ]:
        p(doc, remark, size=10, space_after=3)

    p(doc, '9. Opinion', bold=True, size=12, space_before=6, space_after=4)
    p(doc, 'In our opinion, the Property may be taken as mortgage for the Bank value recommended '
           'in the Property Valuation Report. However, all the remarks made above shall be taken '
           'into consideration and all Legal Documents shall be scrutinized by Legal Expert.\n\n'
           'The said properties have been verified after visiting the site and measuring it as shown '
           'by the Client for the purpose of mortgaging it in {{ bank_name }}.', size=10)

    p(doc, '10. Declaration', bold=True, size=12, space_before=6, space_after=4)
    p(doc, 'This valuation was conducted for the purpose of establishing Fair Market and Distress Values '
           'of the said property for the Client and the Bank for use in mortgaging these properties. '
           'It is not to be used for any other purpose. We certify that our firm is fully authorized to carry '
           'out the valuation work under the prevalent laws and we are fully equipped and competent to carry '
           'out the assignment. No individual in our firm has any financial interest in the said property. '
           'To the best of our knowledge, all matters of a factual nature discussed in this report are true '
           'and correct.', size=10)

    p(doc, '11. Acceptance and Declaration of the Property by the Owner and the Client',
      bold=True, size=12, space_before=6, space_after=4)
    p(doc, 'We are fully responsible for the information and legal document provided to prepare this valuation '
           'report. We are committed, and there is no legal complication. It is not being acquired by '
           'government, and it is not a prohibited area for construction. I also declare that the valuation '
           'is acceptable to me regarding the value of my property and other relevant details.', size=10)

    p(doc, '', space_before=10)
    sig_tbl = doc.add_table(rows=3, cols=3); sig_tbl.style = 'Table Grid'
    set_col_widths(sig_tbl, [2.2, 2.2, 2.2])
    sig_tbl.rows[0].cells[0].text = 'Signature'
    sig_tbl.rows[0].cells[1].text = 'Signature'
    sig_tbl.rows[0].cells[2].text = 'Signature'
    sig_tbl.rows[1].cells[0].text = '{{ owner1_name }}'
    sig_tbl.rows[1].cells[1].text = '{{ owner2_name }}'
    sig_tbl.rows[1].cells[2].text = '{{ client_full_name }}'
    sig_tbl.rows[2].cells[0].text = 'Owner 1'
    sig_tbl.rows[2].cells[1].text = 'Owner 2'
    sig_tbl.rows[2].cells[2].text = 'Client'
    for row in sig_tbl.rows:
        for cell in row.cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)

    p(doc, '12. Certification', bold=True, size=12, space_before=6, space_after=4)
    p(doc, 'I/We hereby certify that Valuation of the Properties as detailed in this report has been '
           'carried out by me/us in strict Compliance with Valuation Guidelines of {{ bank_name }}. '
           'The said property is an acceptable security to bank in all respects.', size=10)

    p(doc, '', space_before=12)
    cert_tbl = doc.add_table(rows=3, cols=2); cert_tbl.style = 'Table Grid'
    set_col_widths(cert_tbl, [3.1, 3.1])
    cert_tbl.rows[0].cells[0].text = 'Site Visited By'
    cert_tbl.rows[0].cells[1].text = 'Checked By'
    cert_tbl.rows[1].cells[0].text = '{{ site_visited_by }}'
    cert_tbl.rows[1].cells[1].text = '{{ certifier_name }}'
    cert_tbl.rows[2].cells[0].text = 'Mobile No: {{ site_visitor_phone }}'
    cert_tbl.rows[2].cells[1].text = 'Mobile No: {{ firm_phone }}'
    for row in cert_tbl.rows:
        for cell in row.cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
    cert_tbl.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    cert_tbl.rows[0].cells[1].paragraphs[0].runs[0].bold = True

    doc.save(os.path.join(DIR, 'proposal_template.docx'))
    print('  ✓ proposal_template.docx')


if __name__ == '__main__':
    print('Building document templates...')
    build_cover()
    build_letterhead()
    build_proposal()
    print('Done.')
