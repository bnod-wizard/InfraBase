"""
Patch proposal_template.docx so triangle tables use dynamic docxtpl row loops.

docxtpl row-loop structure (3 rows required):
  Row A  – {%tr for t in items %}    ← consumed as loop-start tag; no content
  Row B  – {{ t.label }} | {{ t.x }} ← template data row, repeated per item
  Row C  – {%tr endfor %}            ← consumed as loop-end tag; no content

Tables changed (1-indexed, 0-indexed in parentheses):
  Table 24 (idx 23) – measurement triangles  → lm_triangles
  Table 25 (idx 24) – deduction triangles    → ded_triangles
  Table 26 (idx 25) – lalpurja triangles     → lorc_triangles
"""
from copy import deepcopy
from docx import Document


TEMPLATE = 'proposal_template.docx'


def set_cell(cell, text):
    """Replace all content of a cell with plain text, preserving cell XML structure."""
    tc = cell._tc
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    paras = tc.findall(f'{{{ns}}}p')
    for p in paras[1:]:
        tc.remove(p)
    para = cell.paragraphs[0]
    runs = para.runs
    if not runs:
        para.add_run(text)
    else:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ''


def add_row_after(table, after_idx):
    """Insert a deep-copied (cleared) row immediately after after_idx; return it."""
    src_tr = table.rows[after_idx]._tr
    new_tr = deepcopy(src_tr)
    src_tr.addnext(new_tr)
    new_row = table.rows[after_idx + 1]
    for cell in new_row.cells:
        set_cell(cell, '')
    return new_row


DATA_COLS = ['{{ t.label }}', '{{ t.side_a }}', '{{ t.side_b }}',
             '{{ t.side_c }}', '{{ t.semi_perimeter }}', '{{ t.area_sqft }}', '{{ t.aana }}']


def patch_triangle_table(table, loop_var):
    """
    Replace static A/B data rows (rows 2 & 3) with a 3-row docxtpl loop:
      Row 2  – {%tr for t in <loop_var> %}   (tag only, consumed)
      Row 3  – {{ t.label }} | {{ t.side_a }} ...  (data template, repeated)
      Row 4  – {%tr endfor %}               (tag only, consumed)

    Before calling, the table must still have the original A/B rows at indices 2 & 3.
    """
    ncols = len(table.rows[2].cells)

    # Row 2 → for-tag row (NO content, just the tag)
    set_cell(table.rows[2].cells[0], '{' + f'%tr for t in {loop_var} %' + '}')
    for i in range(1, ncols):
        set_cell(table.rows[2].cells[i], '')

    # Row 3 → data template row
    for i, txt in enumerate(DATA_COLS[:ncols]):
        set_cell(table.rows[3].cells[i], txt)

    # Insert Row 4 as endfor row
    endfor_row = add_row_after(table, 3)
    set_cell(endfor_row.cells[0], '{%tr endfor %}')


def patch_lorc_table(table):
    """
    Restructure the lalpurja table from R/A/P/D format to triangle format
    with a 3-row dynamic row loop.

    Before (5 rows):
      0  Header: Location | Plot No. | R | A | P | D | Area in Sq.ft.
      1  Data:   {{ property_location_full }} | {{ plot_no }} | lorc values...
      2  Total Sqft
      3  Total Sqm
      4  Total Aana

    After (8 rows):
      0  Header: Triangle | Side a | Side b | Side c | S | Sqft | Aana
      1  Plot No. | {{ plot_no }} | ...
      2  {%tr for t in lorc_triangles %}   (tag row, consumed)
      3  {{ t.label }} | {{ t.side_a }} ...  (data template, repeated)
      4  {%tr endfor %}                     (tag row, consumed)
      5  Total Sqft  (unchanged)
      6  Total Sqm   (unchanged)
      7  Total Aana  (unchanged)
    """
    ncols = len(table.rows[0].cells)

    # 0. Change header row
    for cell, txt in zip(table.rows[0].cells,
                         ['Triangle', 'Side a', 'Side b', 'Side c',
                          'S=(a+b+c)/2', 'Total area in Sqft.', 'Land in Aana']):
        set_cell(cell, txt)

    # 1. Insert Plot No. row after row 0 (pushes old row 1 to row 2)
    plot_row = add_row_after(table, 0)
    set_cell(plot_row.cells[0], 'Plot No.')
    set_cell(plot_row.cells[1], '{{ plot_no }}')

    # 2. Old row 1 is now row 2 → for-tag row
    set_cell(table.rows[2].cells[0], '{%tr for t in lorc_triangles %}')
    for i in range(1, ncols):
        set_cell(table.rows[2].cells[i], '')

    # 3. Insert data-template row after row 2
    data_row = add_row_after(table, 2)
    for i, txt in enumerate(DATA_COLS[:ncols]):
        set_cell(data_row.cells[i], txt)

    # 4. Insert endfor row after new data row (row 3)
    endfor_row = add_row_after(table, 3)
    set_cell(endfor_row.cells[0], '{%tr endfor %}')

    # Rows 5, 6, 7 (old rows 2, 3, 4) remain unchanged — they reference lorc totals


def main():
    doc = Document(TEMPLATE)

    print('Patching Table 24 (measurement triangles)...')
    patch_triangle_table(doc.tables[23], 'lm_triangles')

    print('Patching Table 25 (deduction triangles)...')
    patch_triangle_table(doc.tables[24], 'ded_triangles')

    print('Patching Table 26 (lalpurja triangles)...')
    patch_triangle_table(doc.tables[25], 'lorc_triangles')

    doc.save(TEMPLATE)
    print('Saved:', TEMPLATE)

    # Verify structure
    doc2 = Document(TEMPLATE)
    for ti in [23, 24, 25]:
        tbl = doc2.tables[ti]
        print(f'\nTable {ti + 1} (idx {ti}) — {len(tbl.rows)} rows:')
        for ri, row in enumerate(tbl.rows):
            print(f'  Row {ri}: {[c.text.strip() for c in row.cells]}')


if __name__ == '__main__':
    main()
