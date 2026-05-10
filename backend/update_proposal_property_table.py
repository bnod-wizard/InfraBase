#!/usr/bin/env python3
"""
Replaces the property schedule table in proposal_template.docx (Table 1)
with the full merged-cell version from letterhead_template.docx.
Run from backend/: python3 update_proposal_property_table.py
"""
import os, copy
from docx import Document

TEMPLATES = os.path.join(os.path.dirname(__file__), 'templates')

def update():
    lh_path   = os.path.join(TEMPLATES, 'letterhead_template.docx')
    prop_path = os.path.join(TEMPLATES, 'proposal_template.docx')

    # Pull the property table from the letterhead (it's the only table there)
    lh = Document(lh_path)
    if not lh.tables:
        raise RuntimeError('letterhead_template.docx has no tables — run create_letterhead_template.py first')
    src_tbl = copy.deepcopy(lh.tables[0]._tbl)

    # Open the proposal and replace Table 1 (the old 8-column property table)
    doc = Document(prop_path)
    if len(doc.tables) < 2:
        raise RuntimeError('proposal_template.docx has fewer than 2 tables — unexpected structure')

    old_tbl = doc.tables[1]._tbl
    old_tbl.getparent().replace(old_tbl, src_tbl)

    doc.save(prop_path)
    print(f'✓  Proposal property table updated  →  {prop_path}')
    print(f'   Tables now: {len(doc.tables)}')

if __name__ == '__main__':
    update()
